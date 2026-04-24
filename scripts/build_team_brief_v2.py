"""Generate v2 of the team-brief docx — deeper, longer, and synced to the 2026-04-19 PDF.

Run:
    python3 scripts/build_team_brief_v2.py
"""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Cm


PROJECT = Path("/Users/dominhhien/Documents/AI/Datathon")
OUT = PROJECT / "Vòng 1 — Phân tích đề & Brainstorm (team brief).docx"


# ---------- helpers ----------

def set_cell_bg(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x2A, 0x44)


def add_para(doc, text: str, bold: bool = False, italic: bool = False, size: int = 11, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p


def add_bullet(doc, text: str, style: str = "List Bullet"):
    doc.add_paragraph(text, style=style)


def add_callout(doc, text: str, bg: str = "FFF4CE"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, bg)
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = False
    run.font.size = Pt(10)


def add_table(doc, headers: list[str], rows: list[list[str]], header_bg: str = "1F2A44"):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        cell.text = ""
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
    for i, r in enumerate(rows, start=1):
        for j, v in enumerate(r):
            cell = tbl.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(v)
            run.font.size = Pt(10)
    return tbl


# ---------- document ----------

doc = Document()

# Margins tighter to pack info
for section in doc.sections:
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)

# Default style
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)

# ---------- TITLE ----------
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("DATATHON 2026 — VÒNG 1")
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = RGBColor(0x1F, 0x2A, 0x44)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("The Gridbreakers · Breaking Business Boundaries")
r.italic = True
r.font.size = Pt(13)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("PHÂN TÍCH ĐỀ BÀI & BRAINSTORM CHIẾN LƯỢC — v2 (chi tiết chuyên sâu)")
r.bold = True
r.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Tài liệu kickoff team — sync với PDF bản 2026-04-19")
r.italic = True
r.font.size = Pt(11)

# ---------- Metadata ----------
add_table(
    doc,
    headers=["Mục", "Thông tin"],
    rows=[
        ["Đội", "The Gridbreakers — 4 thành viên: Hiển (Đội trưởng) · Đồng (Lead Analyst) · Kiên (Insights Lead) · Phúc (ML Engineer). Chi tiết §6. (Re-assign Phúc↔Kiên 2026-04-21.)"],
        ["Ngày phát hành brief (v1)", "2026-04-18"],
        ["Ngày cập nhật brief (v2)", "2026-04-19 (sync theo PDF updated cùng ngày)"],
        ["Deadline nộp Vòng 1", "2026-05-01 — cứng, còn 12 ngày tính từ 04-19"],
        ["Vòng Chung kết", "2026-05-23 tại VinUniversity, Hà Nội (bắt buộc ≥1 thành viên dự)"],
        ["Kaggle competition", "https://www.kaggle.com/competitions/datathon-2026-round-1"],
        ["LaTeX template", "NeurIPS 2025 — https://neurips.cc/Conferences/2025/CallForPapers"],
        ["Vai trò giả định", "Data Scientist tại công ty thời trang TMĐT Việt Nam"],
        ["Dataset", "14 file CSV (~130MB), 10.5 năm dữ liệu (2012-07-04 → 2022-12-31)"],
        ["File kỷ luật nội bộ", "CLAUDE.md (§1–§13) — brain file tổng, đọc trước mọi phiên làm việc"],
    ],
)

# ---------- 1. TL;DR ----------
add_heading(doc, "1. Tóm tắt nhanh (TL;DR)", level=1)
add_para(
    doc,
    "Datathon 2026 Vòng 1 yêu cầu chúng ta đóng vai Data Scientist của 1 công ty thời trang TMĐT VN, "
    "khai thác 14 file CSV (10.5 năm dữ liệu) để biến dữ liệu thành giải pháp kinh doanh. Bài thi 3 phần "
    "với tổng 100 điểm; trọng số áp đảo nằm ở Phần 2 (EDA = 60đ). 3 mục tiêu nghiệp vụ xuyên suốt: "
    "(1) tối ưu tồn kho, (2) lập kế hoạch promo, (3) quản lý logistics theo vùng.",
)

add_table(
    doc,
    headers=["Phần", "Nội dung", "Điểm", "Tỷ trọng", "Ưu tiên đầu tư"],
    rows=[
        ["1", "MCQ — 10 câu × 2đ", "20", "20%", "Trung bình (1 ngày, code 1 notebook)"],
        ["2", "EDA & Data Visualization", "60", "60%", "★★★ CAO NHẤT (5–6 ngày)"],
        ["3", "Sales Forecasting (Kaggle)", "20", "20%", "★★ Cao (3–4 ngày)"],
        ["—", "Tổng", "100", "100%", "—"],
    ],
)

add_callout(
    doc,
    "Insight chiến lược: Phần 2 EDA đáng giá gấp 3 lần Phần 3 Forecasting. Kaggle ranking chỉ chiếm 12/100 điểm tổng "
    "— không cần 'all-in' vào model. Đầu tư mạnh nhất vào EDA storytelling + 4 tầng phân tích (đặc biệt Prescriptive). "
    "Phần 3 cần beat baseline (≥20% MAE) và viết explainability tốt để lấy đủ 8đ báo cáo kỹ thuật.",
)

# ---------- 2. Phân tích đề bài ----------
add_heading(doc, "2. Phân tích đề bài", level=1)

add_heading(doc, "2.1. Bối cảnh kinh doanh", level=2)
add_para(
    doc,
    "PDF §2.3.1: 'Doanh nghiệp cần dự báo nhu cầu (demand) chính xác ở mức chi tiết để tối ưu hoá phân bổ tồn kho, "
    "lập kế hoạch khuyến mãi và quản lý logistics trên toàn quốc.' Dù target model là daily Revenue/COGS, "
    "framing 'demand' mở rộng hơn — nhắc chúng ta rằng báo cáo Prescriptive nên nói về demand, không chỉ tiền.",
)

add_table(
    doc,
    headers=["Mục tiêu nghiệp vụ", "KPI cải thiện", "Dataset liên quan"],
    rows=[
        [
            "1. Tối ưu phân bổ tồn kho",
            "Giảm stockout_days 20–30% cho top SKU; tăng fill_rate; giảm overstock_flag",
            "inventory, order_items, products, sales",
        ],
        [
            "2. Kế hoạch khuyến mãi",
            "Tăng incremental revenue / campaign; giảm cannibalization với khách cũ",
            "promotions, order_items, orders, customers, sales",
        ],
        [
            "3. Logistics toàn quốc",
            "Giảm lead-time (delivery_date − ship_date); phân bổ inventory theo region rising",
            "orders, shipments, geography, inventory",
        ],
    ],
)

add_para(
    doc,
    "→ Mọi insight EDA và feature model phải kết nối ngược về 1 trong 3 mục tiêu này khi viết phần Prescriptive. "
    "Đây là tiêu chí ban giám khảo đánh giá 'business value' (15đ trong Phần 2).",
    italic=True,
)

# ---------- 2.2 Part 1 MCQ ----------
add_heading(doc, "2.2. Phần 1 — MCQ (20 điểm)", level=2)
add_bullet(doc, "10 câu × 2đ = 20đ. KHÔNG trừ điểm câu sai → luôn đoán cả khi không chắc.")
add_bullet(doc, "Mỗi câu chỉ chọn 1 đáp án (radio).")
add_bullet(doc, "Tất cả câu hỏi đều tính trực tiếp từ CSV → phải code thật để verify, không đoán bằng trực giác.")

add_para(doc, "Bảng 10 câu — file nguồn, phương pháp tính, và hint verify:", bold=True)

add_table(
    doc,
    headers=["Q", "Nội dung", "File & join", "Phương pháp tính (pseudocode)"],
    rows=[
        [
            "1",
            "Trung vị số ngày giữa 2 lần mua liên tiếp (inter-order gap) ở khách hàng có >1 đơn",
            "orders.csv",
            "df.sort([customer_id, order_date]); gap = order_date.diff() groupby customer_id; filter >=1 order gap; median(gap).days",
        ],
        [
            "2",
            "Segment có gross margin trung bình cao nhất",
            "products.csv",
            "margin = (price−cogs)/price; groupby segment; mean(margin).idxmax()",
        ],
        [
            "3",
            "Lý do trả hàng phổ biến nhất trong category Streetwear",
            "returns ⨝ products trên product_id",
            "filter products.category=='Streetwear'; value_counts(return_reason).idxmax()",
        ],
        [
            "4",
            "traffic_source có bounce_rate trung bình thấp nhất",
            "web_traffic.csv",
            "groupby traffic_source; mean(bounce_rate).idxmin()",
        ],
        [
            "5",
            "% dòng order_items có promo_id không null",
            "order_items.csv",
            "(promo_id.notna().sum() / len(df)) * 100",
        ],
        [
            "6",
            "age_group có trung bình số đơn / khách cao nhất",
            "orders + customers trên customer_id",
            "orders.groupby(customer_id).size → map age_group → groupby age_group mean(); dropna age_group trước",
        ],
        [
            "7",
            "Region tạo tổng doanh thu cao nhất trong sales_train",
            "order_items ⨝ orders ⨝ geography (orders.zip → geography.zip)",
            "revenue_per_line = quantity × unit_price; groupby region sum(); idxmax(). ⚠ sales.csv không có region — phải dùng order_items.",
        ],
        [
            "8",
            "Payment method phổ biến nhất ở đơn cancelled",
            "orders.csv",
            "filter order_status=='cancelled'; value_counts(payment_method).idxmax()",
        ],
        [
            "9",
            "Size có return rate cao nhất (#returns / #order_items lines)",
            "returns ⨝ products; order_items ⨝ products",
            "num = returns.merge(products).groupby(size).size; den = order_items.merge(products).groupby(size).size; (num/den).idxmax()",
        ],
        [
            "10",
            "Installment plan có giá trị thanh toán trung bình cao nhất",
            "payments.csv",
            "groupby installments; mean(payment_value).idxmax()",
        ],
    ],
)

add_callout(
    doc,
    "Chiến thuật MCQ: 1 người (đề xuất Hiển) viết 02_mcq.ipynb chạy cả 10 câu trong vài giờ. Save kết quả ra "
    "outputs/mcq_answers.json + print full breakdown (mean/count per group) để team verify chéo. KHÔNG submit "
    "trước khi có 2 người đọc lại.",
)

# ---------- 2.3 Part 2 EDA ----------
add_heading(doc, "2.3. Phần 2 — EDA & Data Visualization (60 điểm)", level=2)
add_para(
    doc,
    "Đây là phần ĂN ĐIỂM NHẤT. Không có đáp án đúng duy nhất — BGK chấm khả năng kể chuyện bằng dữ liệu và chiều sâu phân tích. "
    "Bài nộp gồm 2 phần: (1) Visualizations có title/axis/legend rõ; (2) Analysis đi kèm mỗi chart: key findings + business implications + actionable recommendations.",
)

add_para(doc, "Rubric 4 tiêu chí (PDF page 15):", bold=True)
add_table(
    doc,
    headers=["Tiêu chí", "Mô tả", "Điểm tối đa"],
    rows=[
        ["Chất lượng trực quan hoá", "Chart có title/axis/legend đủ, chọn loại phù hợp, thẩm mỹ", "15"],
        ["Chiều sâu phân tích", "Phủ đủ 4 tầng Descriptive → Diagnostic → Predictive → Prescriptive", "25"],
        ["Insight kinh doanh", "Đề xuất cụ thể, định lượng, áp dụng được ngay", "15"],
        ["Tính sáng tạo & kể chuyện", "Góc nhìn độc đáo, narrative xuyên suốt, kết nối đa bảng", "5"],
        ["—", "Tổng", "60"],
    ],
)

add_para(doc, "Chi tiết band điểm (để tự chấm):", bold=True)
add_table(
    doc,
    headers=["Tiêu chí", "Band cao", "Band trung", "Band thấp"],
    rows=[
        ["Trực quan (15đ)", "13–15: tất cả chart đạt chuẩn, chọn loại tối ưu", "8–12: phần lớn đạt, thiếu nhãn", "0–7: thiếu thông tin, khó đọc"],
        ["Chiều sâu (25đ)", "21–25: đạt cả 4 tầng nhất quán", "14–20: 3 tầng, Prescriptive hời hợt", "0–13: chỉ mô tả bề mặt"],
        ["Insight (15đ)", "13–15: đề xuất cụ thể, định lượng, áp dụng ngay", "8–12: đề xuất chung chung", "0–7: thiếu kết nối business"],
        ["Sáng tạo (5đ)", "4–5: góc độc đáo, narrative thuyết phục", "2–3: có sáng tạo chưa nhất quán", "0–1: dự đoán được"],
    ],
)

add_para(doc, "4 tầng phân tích — câu hỏi BGK & chiều sâu kỳ vọng:", bold=True)
add_table(
    doc,
    headers=["Tầng", "Câu hỏi BGK", "Ví dụ phân tích chiều sâu"],
    rows=[
        ["Descriptive", "What happened?", "YoY revenue 2013-2022 + monthly seasonality; top 10 SKU; region split; customer pyramid theo age_group."],
        ["Diagnostic", "Why did it happen?", "COVID 2020-2021: revenue dip vs same-month 2019 (∆%), channel mix shift (mobile spike?); test hypothesis với z-test."],
        ["Predictive", "What is likely?", "STL decompose sales → trend+season+residual; extrapolate; tính dự báo demand SKU-level cho Q1-2023 dựa rolling 365 + cohort retention."],
        ["Prescriptive", "What should we do?", "Shift 20% social → paid_search (ΔROAS định lượng); raise price Premium +5% (margin cao, elasticity < 0.5); pre-stock Tết +14 ngày."],
    ],
)

add_callout(
    doc,
    "Quy tắc vàng: Top team đạt PRESCRIPTIVE NHẤT QUÁN trên nhiều phân tích sẽ được điểm cao nhất. "
    "Mỗi visualization phải có ít nhất 1 dòng 'So what?' + 1 dòng 'Action gì?' + 1 con số định lượng impact (VD: tiết kiệm X triệu/năm, tăng Y% conversion).",
)

# ---------- 2.4 Part 3 Forecasting ----------
add_heading(doc, "2.4. Phần 3 — Sales Forecasting (20 điểm)", level=2)
add_bullet(doc, "Target — Predict Revenue (và COGS) hàng ngày cho 548 ngày 2023-01-01 → 2024-07-01.")
add_bullet(doc, "Train — sales.csv (PDF gọi sales_train.csv) 2012-07-04 → 2022-12-31, 3,833 ngày.")
add_bullet(doc, "Submission — submission.csv 3 cột (Date, Revenue, COGS), thứ tự dòng KHỚP 100% sample_submission.csv. Không reorder.")
add_bullet(doc, "Metrics — MAE, RMSE (càng thấp càng tốt), R² (càng cao càng tốt, lý tưởng ~1).")

add_para(doc, "Quy đổi điểm Phần 3:", bold=True)
add_table(
    doc,
    headers=["Thành phần", "Band", "Mô tả"],
    rows=[
        ["Hiệu suất mô hình (12đ)", "10–12", "Top leaderboard; MAE/RMSE thấp, R² cao"],
        ["", "5–9", "Trung bình; model chạy nhưng chưa tối ưu"],
        ["", "3–4", "Hợp lệ, performance thấp (mức sàn)"],
        ["Báo cáo kỹ thuật (8đ)", "7–8", "Pipeline rõ, temporal CV đúng, SHAP cụ thể, tuân thủ ràng buộc"],
        ["", "4–6", "Pipeline đủ dùng, giải thích định tính, ràng buộc chưa tường minh"],
        ["", "0–3", "Thiếu giải thích, leakage, không tái lập"],
    ],
)

add_callout(
    doc,
    "Điều kiện LOẠI Phần 3 (mất nguyên 20đ): (1) Dùng Revenue/COGS test làm feature; (2) Dùng dữ liệu ngoài 14 CSV; "
    "(3) Không nộp code hoặc kết quả không tái lập. Mitigate ngay từ đầu: anti-leakage checklist + seed=42 + submission "
    "chạy end-to-end trên máy thứ 2 trước khi nộp.",
    bg="FFE0E0",
)

# ---------- 3. Dataset ----------
add_heading(doc, "3. Dataset tour — 14 CSV chia 4 lớp", level=1)
add_para(
    doc,
    "PDF §1.1 tổ chức 14 file thành 4 lớp logic. PDF intro ghi '15 file CSV' nhưng Table 1 chỉ liệt kê 14 — "
    "trên disk thực tế là 14 (không có inventory_enhanced.csv; inventory.csv hiện tại đã có đủ cột enhanced).",
)

add_table(
    doc,
    headers=["Lớp", "Vai trò", "File", "Đặc điểm khi dùng"],
    rows=[
        ["Master", "Tham chiếu, ít thay đổi", "products, customers, promotions, geography", "Slowly-changing dimensions; dùng để dim, filter, groupby."],
        ["Transaction", "Sự kiện thô event-level", "orders, order_items, payments, shipments, returns, reviews", "Khối lượng lớn (~647K orders, 715K items). Gốc của mọi metric daily."],
        ["Analytical", "Aggregate pre-compute", "sales (train), sample_submission (test format)", "Daily Revenue + COGS — chính là TARGET của Phần 3."],
        ["Operational", "Vận hành nội bộ", "inventory (monthly snapshot), web_traffic (daily)", "Inventory end-of-month; web_traffic không có cho test period."],
    ],
)

add_para(doc, "Quan hệ giữa các bảng (cardinality):", bold=True)
add_table(
    doc,
    headers=["Quan hệ", "Cardinality"],
    rows=[
        ["orders ↔ payments", "1 : 1"],
        ["orders ↔ shipments", "1 : 0 hoặc 1 (chỉ status shipped/delivered/returned)"],
        ["orders ↔ returns", "1 : 0 hoặc nhiều"],
        ["orders ↔ reviews", "1 : 0 hoặc nhiều (~20% delivered orders có review)"],
        ["order_items ↔ promotions", "nhiều : 0 hoặc 1 (qua promo_id, promo_id_2)"],
        ["products ↔ inventory", "1 : nhiều (1 dòng/sản phẩm/tháng)"],
    ],
)

# ---------- 3.1 Quick reference ----------
add_heading(doc, "3.1. Quick-reference — cột đắt giá mỗi CSV", level=2)
add_para(
    doc,
    "Bảng dưới để join 2+ bảng mà không phải mở dataset_description.docx. PK = primary key, FK = foreign key. "
    "Cột in đậm ở notebook = thường dùng làm feature hoặc groupby.",
)

add_table(
    doc,
    headers=["File", "PK / FK", "Cột đắt giá cho phân tích"],
    rows=[
        ["products", "PK product_id", "category (4), segment (8), size (S/M/L/XL), color, price, cogs"],
        ["customers", "PK customer_id, FK zip", "signup_date, age_group (nullable), gender (nullable), acquisition_channel (nullable)"],
        ["geography", "PK zip", "region (East/Central/West), city (42), district"],
        ["promotions", "PK promo_id", "promo_type (percentage/fixed), discount_value, start_date, end_date, applicable_category, stackable_flag, min_order_value"],
        ["orders", "PK order_id, FK customer_id, zip", "order_date, order_status (6), payment_method (5), device_type (3), order_source (6)"],
        ["order_items", "PK (order_id, product_id); FK promo_id, promo_id_2", "quantity, unit_price (ĐÃ giảm giá!), discount_amount"],
        ["payments", "FK order_id (1:1)", "payment_method, payment_value, installments (1/3/6/12)"],
        ["shipments", "FK order_id (1:0/1)", "ship_date, delivery_date, shipping_fee (0 = free)"],
        ["returns", "PK return_id; FK order_id, product_id", "return_date, return_reason (5), return_quantity, refund_amount"],
        ["reviews", "PK review_id; FK order_id, product_id, customer_id", "review_date, rating (1-5), review_title"],
        ["inventory", "PK (snapshot_date, product_id)", "stock_on_hand, stockout_days, days_of_supply, fill_rate, stockout_flag, overstock_flag, reorder_flag, sell_through_rate (+category, segment denorm)"],
        ["web_traffic", "PK date (từ 2013-01-01)", "sessions, unique_visitors, page_views, bounce_rate, avg_session_duration_sec, traffic_source (6). KHÔNG có conversion_rate → tự tính orders/sessions."],
        ["sales", "PK Date (2012-07-04 → 2022-12-31, 3833 dòng)", "Revenue, COGS — ★ TARGET train. PDF dùng cả tên sales.csv (Table 1) lẫn sales_train.csv (Split). File thật trên disk: sales.csv."],
        ["sample_submission", "PK Date (2023-01-01 → 2024-07-01, 548 dòng)", "Revenue, COGS dummy — ★ format test. File test thật (sales_test.csv) không public."],
    ],
)

# ---------- 3.2 Data gotchas ----------
add_heading(doc, "3.2. 7 cái bẫy dữ liệu PHẢI biết", level=2)

add_para(doc, "Đã gặp hoặc dự kiến gặp khi code — ghi ra để không ai lặp lại:", italic=True)
traps = [
    ("order_items.unit_price là GIÁ SAU GIẢM", "Không phải giá list. Muốn ra giá gốc: (qty × unit_price + discount_amount) / qty. Nếu quên, tất cả phân tích pricing / gross margin sẽ sai."),
    ("PDF dùng 2 tên cho cùng file train", "Table 1 gọi sales.csv; Split table và Q7 gọi sales_train.csv. File thật trên disk: sales.csv. Trích dẫn trong báo cáo nên thống nhất 1 tên (nên dùng 'sales.csv (Train split)')."),
    ("inventory_enhanced.csv KHÔNG tồn tại", "PDF cũ liệt kê file #14, PDF 2026-04-19 đã xoá. inventory.csv hiện tại đã có enhanced columns (stockout_flag, overstock_flag, reorder_flag, sell_through_rate). Không đi tìm file riêng."),
    ("web_traffic.csv chỉ có đến 2022-12-31", "Test period 2023-01-01 → 2024-07-01 KHÔNG có traffic. 3 option xử lý: (a) forecast traffic riêng, (b) rolling mean 2022 proxy, (c) bỏ traffic feature. Decide trước khi build FE pipeline."),
    ("shipments chỉ tồn tại cho 3 status", "Chỉ shipped/delivered/returned. Orders status cancelled/paid/created không có dòng shipment → left join sẽ có NaN, đừng coi là data quality issue."),
    ("web_traffic.csv bắt đầu 2013-01-01, sales.csv bắt đầu 2012-07-04", "6 tháng đầu của sales KHÔNG có dữ liệu traffic. Dùng inner join sẽ mất 180 ngày, hoặc chấp nhận NaN và cho model học."),
    ("inventory.csv snapshot end-of-month, không có daily", "Phải forward-fill sang daily khi join với sales. Không có inventory cho 2023+ (giống web_traffic) → cần forward-fill 2022-12-31 snapshot hoặc drop feature."),
]
rows_traps = [[f"#{i+1}. {t[0]}", t[1]] for i, t in enumerate(traps)]
add_table(doc, headers=["Cạm bẫy", "Chi tiết & cách xử"], rows=rows_traps)

# ---------- 4. Brainstorm ----------
add_heading(doc, "4. Brainstorm — ý tưởng tổng hợp & chiến lược", level=1)

add_callout(
    doc,
    "Bản đồ §4 (đọc theo luồng):\n"
    "  §4.1 Narrative (khung story) ─▶ quyết định chart nào vào báo cáo.\n"
    "  §4.2 12 ý tưởng EDA ─▶ cung cấp insight cho §4.4 (Prescriptive) + feature cho §4.3 (Model).\n"
    "  §4.3 Feature blocks ─▶ input cho notebook 05/06/07 (Kiên). Mỗi block link về ý EDA nguồn.\n"
    "  §4.4 Prescriptive KPI ─▶ ending của story §4.1; phục vụ 3 business goal §2.1; con số 'Dự kiến impact' do §4.3 forecast cung cấp.\n"
    "  §4.5 Coverage matrix ─▶ audit: mỗi ý §4.2 có phủ đủ 14 CSV không.",
    bg="E8F1FF",
)

add_heading(doc, "4.1. Narrative xuyên suốt (đề xuất)", level=2)
add_para(doc, "→ Phụ trách: Kiên (Insights Lead). Tham chiếu: §4.2 (chart), §4.4 (ending = action).", italic=True, size=10)
add_callout(
    doc,
    'Tagline báo cáo: "Từ DATA đến DECISION — Hành trình tăng trưởng 10 năm của một thương hiệu thời trang Việt Nam." '
    "Cấu trúc story 5 nút:\n"
    "  (1) Quá khứ 2013-2019 — tăng trưởng  → chart từ §4.2 ý #4 (regional growth) + #11 (segment mix).\n"
    "  (2) Cú sốc COVID 2020-2021  → chart từ §4.2 ý #10 (seasonality + event overlay).\n"
    "  (3) Phục hồi 2022  → chart từ §4.2 ý #2 (promo ROI) + #5 (inventory-sales mismatch).\n"
    "  (4) Dự báo 2023-2024  → chart từ §4.3 (pred vs actual + SHAP summary).\n"
    "  (5) Khuyến nghị hành động cho 3 business goals  → toàn bộ §4.4.",
)

add_para(
    doc,
    "Vì báo cáo ≤4 trang, chỉ nên giữ tối đa 5-6 chart 'đắt giá nhất' cho EDA + 2 chart cho Forecasting "
    "(pred vs actual trên val, SHAP summary). Appendix có thể chứa chart phụ. "
    "Kỷ luật chọn chart: mỗi chart phải gắn với 1 nút trong story arc trên + 1 hàng trong §4.4 (nếu không có action → cắt).",
)

add_callout(
    doc,
    "4 story arc dự phòng (thematic lens — dùng khi viết notebook 04 cần chia section theo chủ đề, "
    "hoặc khi narrative 5-nút không fit slot chart):\n"
    "  • Story A — 'Doanh thu không chỉ nằm ở traffic': consume ý #6 (web funnel) + §4.4 'Channel-mix budget'.\n"
    "  • Story B — 'Khuyến mãi là con dao hai lưỡi': consume ý #2 (Promo ROI) + #9 (cannibalization) + §4.4 'Promo scheduling'.\n"
    "  • Story C — 'Tồn kho & size ảnh hưởng trải nghiệm': consume ý #3 (return × size) + #5 (inventory mismatch) + §4.4 'Size-guide fix' + 'Inventory reorder'.\n"
    "  • Story D — 'Nhóm khách hàng mục tiêu': consume ý #1 (cohort retention) + #11 (segment mix) + §4.4 'Customer retention' + 'Pricing & discount'.\n"
    "Narrative 5-nút (Quá khứ → COVID → Phục hồi → Dự báo → Khuyến nghị) vẫn là PRIMARY story cho PDF ≤4 trang. "
    "4 story arc trên là lens bổ trợ cho notebook section, không thay thế narrative chính.",
    bg="E6F4EA",
)

# ---------- 4.2 EDA ideas — DEEPENED (TTM framework) ----------
add_heading(doc, "4.2. 12 ý tưởng EDA — Framework TTM (Pain → Truth → Tension → Motivation → Insight → Action)", level=2)
add_para(doc, "→ Phụ trách: Đồng (ý D/Di) + Kiên (ý P/Pr). Mỗi ý chỉ rõ section nào của §4 sẽ 'tiêu thụ' kết quả.", italic=True, size=10)
add_callout(
    doc,
    "Framework TTM — chuẩn 'lọt top' cho mỗi ý EDA (bắt buộc đủ 6 layer):\n"
    "  • Pain — business pain point cụ thể (ai đau, đau ở đâu, đau bao nhiêu tiền).\n"
    "  • Truth — fact từ data sau phân tích (con số + chart reference, observable, verifiable).\n"
    "  • Tension — mâu thuẫn ngầm customer không nói ra, là chỗ business đang hiểu sai.\n"
    "  • Motivation — động lực sâu behind customer behavior (lý do thật họ hành xử thế).\n"
    "  • Insight (the one line) — câu chốt ≤20 từ; BGK đọc 1 lần là nhớ. Insight phải 'đắt hơn truth'.\n"
    "  • Action — hành động business team làm ngay → liên kết §4.4 + 1 business goal §2.1.\n"
    "Nối tiếp plan cũ: mỗi ý có 2 link explicit → §4.1 nút (N) + §4.4 hàng 'X'. Block 'Execution note' giữ H/T/A cũ cho người phân tích chạy code.\n"
    "Icon 🔄 = counter-intuitive takeaway (nhắm điểm 'Tính sáng tạo & kể chuyện' §2.3 rubric).",
    bg="E8F1FF",
)
add_para(
    doc,
    "Lưu ý về template: chart-level annotation trong notebook 03/04 dùng O-E-I-A (Observation → Evidence → Interpretation → Action) cho gọn. "
    "Insight-level callout trong PDF báo cáo PHẢI dùng TTM — TTM supersedes O-E-I-A cho báo cáo ≤4 trang.",
    italic=True,
    size=10,
)

eda_ideas = [
    (
        "1. Customer lifecycle & cohort retention",
        "customers ⨝ orders (customer_id)",
        "Di + Pr",
        "Đồng",
        False,
        (
            "CAC cao nhưng churn 90-day ở social_media lớn → marketing đốt tiền không hoàn vốn.",
            "Cohort 90-day retention paid_search ≈ X% vs social_media ≈ Y% (cohort heatmap count distinct customer_id quay lại theo tháng).",
            "Marketing nghĩ social là 'brand awareness' nên churn OK; data cho thấy social user không bao giờ quay lại mua lần 2.",
            "Paid_search đến với nhu cầu cụ thể ('tôi cần áo mùa thu') → loyal vì problem-solution fit. Social đến vì giải trí, không cần sản phẩm.",
            "Social_media mua 1 lượt — paid_search mới mua lần 2.",
            "Email re-engage cohort social sau 90 ngày; re-allocate budget sang paid_search nếu retention gap >2×.",
        ),
        "H: paid_search retention 90-day > social_media. T: cohort heatmap = count distinct customer_id quay lại tháng M+k / tại M, split theo acquisition_channel. A: tăng CPA bid paid_search nếu H đúng; nếu sai → social cần retention campaign.",
        "→ §4.1 nút (1) Quá khứ tăng trưởng  ·  → §4.4 hàng 'Customer retention'  ·  → §4.3 block 'Customer/Return signals (lag_365)'",
    ),
    (
        "2. Promo ROI deep-dive (50 campaign / 10.5 năm — sparse)",
        "promotions ⨝ order_items ⨝ orders ⨝ sales",
        "Di + Pr",
        "Kiên",
        True,
        (
            "50 campaign tốn discount cost 10.5 năm — không rõ campaign nào thật sự incremental vs pull-forward demand.",
            "Top 5 campaign Δrevenue vs same-week-prior-year ≈ +X%; bottom 5 ≈ −Y% (cannibalize baseline).",
            "Marketing nghĩ mỗi discount đều 'kích cầu'; số liệu cho thấy campaign dài chỉ pull-forward, không tạo demand mới.",
            "Khách mua khi promo không vì 'giá rẻ' mà vì deadline (FOMO). Campaign >7 ngày không có urgency → cannibalize baseline.",
            "Promo không tạo demand — promo chỉ dịch chuyển thời điểm mua. Chỉ scarcity mới incremental.",
            "Rút promo window ≤3 ngày; target segment chưa mua 90 ngày thay vì blast-all.",
        ),
        "H: promo percentage có incremental revenue > fixed amount. T: cho mỗi campaign tính Δrevenue trong promo window vs same-week prior year, split promo_type. A: Top 5 case study; Bottom 5 rút kinh nghiệm; quantify $X incremental/năm nếu scale top pattern.",
        "→ §4.1 nút (3) Phục hồi  ·  → §4.4 hàng 'Promo scheduling'  ·  → §4.3 block 'Promo indicators'  ·  Liên quan MCQ Q5",
    ),
    (
        "3. Return reason × size × category heatmap",
        "returns ⨝ order_items ⨝ products",
        "Di + Pr",
        "Đồng",
        False,
        (
            "Mỗi return = phí ship 2 chiều + refund + hàng second-hand khó bán lại → margin ăn mòn.",
            "wrong_size chiếm ≥40% returns ở Streetwear; tập trung size S/XL (extremes) — pivot heatmap (category, size, reason).",
            "Ops nghĩ khách 'đổi ý'; thực chất size chart lệch so với body thật của Gen-Z VN.",
            "Khách Gen-Z mua online theo size số (không thử); lệch 1 size = buộc phải trả, không phải 'whim'.",
            "Return không phải 'đổi ý' — đó là size chart đang nói dối khách.",
            "Fix size guide top 3 SKU wrong_size cao nhất; thêm tag 'chạy đúng/to/nhỏ' từ review users.",
        ),
        "H: wrong_size chiếm ≥40% returns ở Streetwear, tập trung S/XL. T: pivot count(return_id) trên (category, size, reason) + chi-square. A: fix size guide top 3 SKU; refund_saved = N_wrong_size × avg_refund × 30% giảm.",
        "→ §4.1 nút (3) Phục hồi  ·  → §4.4 hàng 'Size-guide fix'  ·  → §4.3 block 'Customer/Return signals (lag_365)'  ·  Liên quan MCQ Q3",
    ),
    (
        "4. Regional growth map (East/Central/West)",
        "order_items ⨝ orders ⨝ geography",
        "D + P",
        "Đồng",
        False,
        (
            "Warehouse đặt cố định (giả định HN/miền Bắc) trong khi demand East (HCM) tăng nhanh → lead-time East dài, chi phí reverse logistics cao.",
            "East YoY revenue growth 2020-2022 ≈ +X%/năm, West ≈ +Y%/năm (choropleth region × year, linear fit YoY, CI 95%).",
            "Supply chain nghĩ 'các vùng mua như nhau'; gravity demand đã dịch về East từ 2020.",
            "HCM đô thị hoá nhanh, Gen-Z online-first; miền Bắc-Trung còn offline-first.",
            "Demand đã chuyển về East, nhưng warehouse vẫn ở đâu đó giữa BC 2018.",
            "Mở warehouse vệ tinh East → giảm lead-time 5 → 3 ngày.",
        ),
        "H: East tăng trưởng YoY nhanh nhất 2020-2022. T: choropleth region × year; linear fit. A: pre-stock inventory East; cross-check với shipments.lead_time per region.",
        "→ §4.1 nút (1) Quá khứ tăng trưởng  ·  → §4.4 hàng 'Regional logistics'  ·  → §4.3 block 'Rolling features' (region-level optional)  ·  Liên quan MCQ Q7",
    ),
    (
        "5. Inventory–sales mismatch (lost revenue)",
        "inventory ⨝ order_items ⨝ orders ⨝ sales",
        "Di + Pr",
        "Đồng",
        False,
        (
            "Stockout = lost revenue; overstock = vốn chết + rủi ro giảm giá cuối mùa.",
            "~15% peak-demand SKU-day có stockout_flag=1 trong Q4 (mùa lễ); lost_revenue = avg_daily_qty × stockout_days × unit_price.",
            "Supply chain báo 'đủ hàng' theo tổng; SKU-level cho thấy bestseller hết trong khi slow-mover dư.",
            "Buyer tính theo tổng doanh thu, không theo SKU profile → không phân bổ đủ cho bestseller.",
            "Out-of-stock không phải vấn đề supply — là vấn đề phân bổ. Bestseller thiếu trong khi slow-movers dư.",
            "Reorder threshold = days_of_supply × 1.3 cho top 50 SKU (80% revenue); reclaim $X revenue/năm.",
        ),
        "H: ≥15% peak-demand SKU-day rơi vào stockout_flag=1 Q4. T: forward-fill inventory daily × order_items; lost_revenue. A: reorder threshold top 50 SKU.",
        "→ §4.1 nút (3) Phục hồi  ·  → §4.4 hàng 'Inventory reorder'  ·  → §4.3 block 'Inventory signals'",
    ),
    (
        "6. Web-traffic funnel (orders-per-session proxy)",
        "web_traffic + orders (attribute theo source)",
        "D + P",
        "Đồng",
        True,
        (
            "Marketing budget phân đều 6 channel không biết channel nào thật sự convert → ROAS ảo.",
            "paid_search bounce cao nhưng (n_orders/sessions) × AOV ≈ cao nhất; email bounce thấp nhưng AOV thấp.",
            "Marketing coi 'bounce thấp = tốt'; bounce cao có thể là user biết mình cần gì, click-and-buy.",
            "User paid_search đến với intent rõ (có search query); user email đến vì tò mò newsletter.",
            "Bounce cao không phải user hỏng — có thể là user biết mình cần gì và mua thẳng.",
            "Shift 20% budget social → paid_search; track conv × AOV thay vì bounce.",
        ),
        "H: paid_search bounce cao nhưng conv cao nhất vì intent; email bounce thấp nhưng AOV thấp. T: (conv_proxy = n_orders_attributed / sessions) × AOV, plot 2D. A: optimize bid → channel có conv × AOV cao. ⚠ Traffic chỉ 2013-2022, không có test period.",
        "→ §4.1 nút (4) Dự báo  ·  → §4.4 hàng 'Channel-mix budget'  ·  → §4.3 block 'Traffic features' (cần decide option a/b/c — §9)  ·  Liên quan MCQ Q4",
    ),
    (
        "7. Payment & device behaviour",
        "orders ⨝ payments",
        "D + Di",
        "Đồng",
        True,
        (
            "COD cancel ~30% → mất ship đi + warehouse picking cost ăn vào margin.",
            "COD cancel rate ≈ 30%, credit_card ≈ 5% (groupby(payment_method) % status=='cancelled').",
            "Shop nghĩ COD là service 'build trust với khách mới'; khách COD đã không trust ngay từ đầu.",
            "Chọn COD vì không tin shop → chờ 3-5 ngày ship khuếch đại nghi ngờ → cancel.",
            "COD không phải payment method — đó là niềm tin còn thiếu.",
            "Voucher 2% prepaid cho COD khách cũ; ship express COD <24h để rút window nghi ngờ.",
        ),
        "H: COD có cancel rate cao nhất ~30%. T: groupby(payment_method) % cancelled + chi-square. A: voucher 2% cho COD; giảm cancel 5% → tiết kiệm $X.",
        "→ §4.1 nút (3) Phục hồi  ·  → §4.4 hàng 'COD optimization'  ·  → §4.3 block 'Order aggregates (pct_cod_d, pct_cancelled_d)'  ·  Liên quan MCQ Q8/Q10",
    ),
    (
        "8. Review → return correlation",
        "reviews ⨝ orders ⨝ returns ⨝ products",
        "Di + P",
        "Kiên",
        False,
        (
            "Review xấu damage lên cả product listing, không chỉ 1 đơn — ảnh hưởng lâu dài.",
            "Orders có rating ≤2 có return rate cao gấp ~3× orders rating ≥4 (2×2 contingency + fisher exact).",
            "CSKH chỉ phản hồi sau khi review public — đã muộn, reputation damage đã xảy ra.",
            "Khách để review thấp không vì ghét — vì muốn được nghe. Phản hồi kịp thời → đa số xoá review.",
            "Review thấp là SOS, không phải feedback. Nghe kịp thì cứu được cả đơn lẫn danh tiếng.",
            "Trigger CSKH reach-out ≤48h sau review ≤2; giảm return 20% nếu H đúng.",
        ),
        "H: orders rating ≤2 có return rate 3× orders rating ≥4. T: 2×2 contingency (rating ≤2 vs ≥3) × (returned vs not) + fisher exact. A: CSKH trigger post-review ≤48h.",
        "→ §4.1 nút (3) Phục hồi  ·  → §4.4 hàng 'Size-guide fix' (bổ sung CSKH post-review trigger)",
    ),
    (
        "9. Discount cannibalization",
        "promotions ⨝ order_items ⨝ orders ⨝ customers",
        "Di + Pr",
        "Kiên",
        True,
        (
            "Discount cost $X/năm không biết bao nhiêu % là 'giảm giá cho khách vốn đã sẵn sàng mua full price'.",
            ">60% revenue promo đến từ returning customers (đã mua ≤30 ngày trước) — split pre-promo vs during-promo × flag first-time/returning.",
            "CMO nghĩ promo là 'acquisition tool'; thực chất là 'retention discount' không ai gọi tên.",
            "Returning customer sẽ mua full price nếu không có promo; khi có promo, họ đợi rồi mua. Promo đang train khách đợi sale.",
            "Promo đại trà không acquire khách mới — mà dạy khách cũ đợi sale.",
            "Target segment chưa mua 90 ngày (behavioural trigger); không blast-all.",
        ),
        "H: >60% revenue promo đến từ returning customers → cannibalize. T: split pre vs during promo, flag first-time vs returning, % new customer acquired trong promo. A: target segment dormant 90 ngày.",
        "→ §4.1 nút (3) Phục hồi  ·  → §4.4 hàng 'Promo scheduling' (bổ sung targeting)  ·  → §4.3 block 'Promo indicators'  ·  Liên quan MCQ Q5",
    ),
    (
        "10. Seasonality: Tết / 11.11 / 12.12 / Black Friday",
        "sales + promotions + self-built VN calendar",
        "P + Pr",
        "Kiên",
        True,
        (
            "Promo budget lãng phí khi trùng peak mùa vụ (demand tự peak không cần boost); pre-stock cũng lệch.",
            "Revenue tuần Tết +0→+7 thấp hơn baseline ≈ 30%; 11.11 peak ≈ +200% baseline tự phát (STL decompose + event overlay).",
            "Merch team chạy promo Tết và 11.11 vì 'đó là mùa sale'; Tết promo mất cost, 11.11 promo cannibalize margin.",
            "Khách VN mua theo pattern văn hoá (Tết ở với gia đình; 11.11 là shopping festival của họ) — không cần shop kích.",
            "Mùa vụ không cần promo — promo cần đánh ngoài mùa.",
            "KHÔNG promo tuần Tết +7; pre-stock +14 ngày trước Tết; 'pre-11.11' 1 tuần bắt early shopper thay vì 'all-in 11.11'.",
        ),
        "H: revenue tuần Tết +0→+7 thấp hơn baseline 30% (bận lễ, logistics đóng). T: STL decompose + event overlay; boxplot trong vs ngoài window. A: tránh launch promo Tết; pre-stock 14 ngày trước Tết; feed holiday_flag vào model Phần 3.",
        "→ §4.1 nút (2) COVID + (4) Dự báo  ·  → §4.4 hàng 'Promo scheduling' + 'Inventory reorder'  ·  → §4.3 block 'Calendar features (self-built)'",
    ),
    (
        "11. Gross-margin × segment mix",
        "products ⨝ order_items ⨝ orders",
        "D + Pr",
        "Đồng",
        True,
        (
            "Segment Premium margin 60% nhưng revenue share giảm 2021-2022 → mất vũ khí lợi nhuận.",
            "Premium volume 2021 ≈ X đơn, 2022 ≈ Y đơn (−Z%); Activewear volume +W% (Pareto 80/20 + trend YoY per segment).",
            "Merch nghĩ 'khách trẻ chỉ muốn Activewear giá rẻ'; thực chất Premium bị bỏ quên, không phải không còn khách.",
            "Khách ngân sách Premium vẫn tồn tại, nhưng không được target ads, không thấy collection mới → chuyển brand khác.",
            "Premium không chết — Premium bị bỏ đói marketing.",
            "Bundle Premium + Activewear cross-category; raise price Premium +5% khi elasticity ước tính <0.5.",
        ),
        "H: segment Premium margin ~60% nhưng volume giảm 2021-2022. T: margin % theo segment + revenue trend YoY; Pareto 80/20. A: bundle Premium + Activewear; raise price +5% khi elasticity <0.5.",
        "→ §4.1 nút (1) Quá khứ tăng trưởng  ·  → §4.4 hàng 'Pricing & discount'  ·  Liên quan MCQ Q2",
    ),
    (
        "12. Shipping fee vs return rate",
        "shipments ⨝ orders ⨝ returns",
        "Di + Pr",
        "Kiên",
        True,
        (
            "Free-shipping tốn ship cost, tưởng boost doanh thu nhưng return rate ăn hết margin.",
            "Free-ship orders có return rate cao hơn paid-ship ≈ 15% (odds ratio >1.2; 2×2 shipping_fee==0 vs >0 × returned vs not).",
            "Growth team đẩy free-ship vì 'conversion lift'; return rate ăn hết margin gain.",
            "Free-ship khiến khách 'thử miễn phí, rủi ro = 0' → mua nhiều size/color rồi trả phần thừa (try-before-buy abuse).",
            "Free-ship không kéo khách — free-ship tạo fitting room miễn phí.",
            "min_order_value cho free-ship; cap số SKU free-ship return / customer / tháng.",
        ),
        "H: free-ship orders return rate cao hơn 15%. T: 2×2 table (shipping_fee==0 vs >0) × (returned vs not); odds ratio + CI. A: min_order_value cho free-ship.",
        "→ §4.1 nút (3) Phục hồi  ·  → §4.4 hàng 'Regional logistics' (bổ sung min_order_value rule)",
    ),
]
ttm_labels = ("Pain", "Truth", "Tension", "Motivation", "Insight", "Action")
for title, datasets, layer, owner, counter_intuitive, ttm, execution, links in eda_ideas:
    p = doc.add_paragraph()
    icon = "🔄 " if counter_intuitive else ""
    r = p.add_run(f"★ {icon}{title}")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x1F, 0x2A, 0x44)
    add_para(doc, f"  · Datasets: {datasets}  ·  Tầng: {layer}  ·  Lead: {owner}", italic=True, size=10)
    for label, content in zip(ttm_labels, ttm):
        p_ttm = doc.add_paragraph()
        r_lab = p_ttm.add_run(f"  {label}: ")
        r_lab.bold = True
        r_lab.font.size = Pt(10)
        r_val = p_ttm.add_run(content)
        r_val.font.size = Pt(10)
        if label == "Insight":
            r_val.bold = True
            r_val.font.color.rgb = RGBColor(0xB7, 0x1C, 0x1C)
    p_links = doc.add_paragraph()
    r_links = p_links.add_run(f"  {links}")
    r_links.italic = True
    r_links.font.size = Pt(10)
    r_links.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
    p_exec = doc.add_paragraph()
    r_exec_lab = p_exec.add_run("  Execution note: ")
    r_exec_lab.bold = True
    r_exec_lab.italic = True
    r_exec_lab.font.size = Pt(9)
    r_exec_lab.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    r_exec_val = p_exec.add_run(execution)
    r_exec_val.italic = True
    r_exec_val.font.size = Pt(9)
    r_exec_val.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

add_callout(
    doc,
    "★ Counter-intuitive takeaway (icon 🔄): 6 ý #2, #6, #7, #10, #11, #12 có insight đảo ngược trực giác — "
    "dùng để 'đánh' điểm Tính sáng tạo & Kể chuyện (5đ, §2.3 rubric). "
    "Mỗi ý tag Pr phải kết nối 1 trong 3 business goals (inventory / promo / logistics — §2.1 / CLAUDE.md §1) khi viết insight báo cáo. "
    "Chart nào không dẫn đến action cụ thể nên cắt. "
    "Luật nhất quán: nếu §4.2 tạo ra insight X mà §4.3 không có feature tương ứng + §4.4 không có hàng action → một trong hai bên đang miss; báo ngay cho cả team.",
)

# ---------- 4.3 Model ideas — DEEPENED ----------
add_heading(doc, "4.3. Model Phần 3 — feature block chi tiết", level=2)
add_para(doc, "→ Phụ trách: Phúc (ML Engineer). Cột '→ EDA §4.2' chỉ rõ ý tưởng EDA nào validate/cung cấp cho block này.", italic=True, size=10)
add_para(
    doc,
    "Target: predict Revenue + COGS hàng ngày 548 ngày 2023-01-01 → 2024-07-01. "
    "Bảng dưới liệt kê từng khối feature, nguồn CSV, giá trị lag/window cụ thể, ràng buộc anti-leakage, và liên kết ngược về §4.2.",
)

add_table(
    doc,
    headers=["Feature block", "Cột / Biến tạo", "CSV nguồn & ràng buộc", "→ EDA §4.2 nguồn"],
    rows=[
        [
            "Baseline seasonal × growth",
            "base_level × growth^years_ahead × seasonal(month, day)",
            "sales.csv duy nhất. Benchmark tối thiểu; model nâng cao phải beat ≥20% MAE.",
            "(nền) ý #10 Seasonality + §4.1 story arc",
        ],
        [
            "Lag features (Revenue, COGS)",
            "lag_7, lag_14, lag_28, lag_365; diff_7, diff_365",
            "sales.csv. Test xa 548 ngày → phải recursive forecast (predict step-by-step) hoặc chỉ dùng lag ≥548 (lag_548, lag_730).",
            "ý #10 Seasonality (lý do chọn lag_365)",
        ],
        [
            "Rolling features",
            "rolling_mean_7/30/90, rolling_std_7/30, rolling_min_max_30",
            "sales.csv. Rolling KHÔNG bao gồm day t (shift 1). Test 2023+ → phải recursive hoặc dùng rolling dựa trên lag ≥548.",
            "ý #4 Regional growth + ý #11 Segment mix (trend mượt)",
        ],
        [
            "Calendar features (SELF-BUILT)",
            "day_of_week, month, quarter, is_month_end, is_tet (lunar), is_quoc_khanh, is_black_friday, is_1111, is_1212",
            "⚠ KHÔNG import Prophet holidays (external data → LOẠI). Tự build DataFrame VN holiday 2012-2024 từ list tĩnh.",
            "ý #10 Seasonality (VN event list)",
        ],
        [
            "Order aggregates (historical)",
            "n_orders_d, n_unique_customers_d, pct_cancelled_d, pct_cod_d, pct_mobile_d, avg_AOV_d",
            "orders group by order_date. Các agg này KHÔNG có cho test 2023+ → dùng LAG (ví dụ lag_365) hoặc forecast riêng.",
            "ý #7 Payment/device behaviour",
        ],
        [
            "Order-item aggregates",
            "n_lines_d, n_units_sold_d, pct_promo_d, avg_discount_pct_d, n_unique_SKU_d",
            "order_items group by order_date. Tương tự: chỉ dùng lag cho test period.",
            "ý #2 Promo ROI + ý #9 Cannibalization",
        ],
        [
            "Traffic features",
            "sessions_d, bounce_rate_d, page_views_d, pct_paid_search_d, pct_social_d, avg_duration_d, engagement_index (= avg_session_duration × (1 − bounce_rate)), traffic_momentum (= sessions / sessions.rolling(7).mean())",
            "web_traffic daily. Chỉ có 2013-2022, không có test → Option (a) forecast riêng / (b) rolling 2022 proxy / (c) drop. Decide Phase 6.",
            "ý #6 Web funnel (cần output conv_proxy)",
        ],
        [
            "Promo indicators",
            "active_promo_today, n_active_promos, days_since_last_promo, days_until_next_promotion, days_since_promotion, promo_type_one_hot",
            "promotions expand daily. 50 campaign đã biết lịch 2013-2022. Test 2023+ → dùng historic seasonality như proxy hoặc gán all-zero (conservative).",
            "ý #2 Promo ROI + ý #9 Cannibalization",
        ],
        [
            "Inventory signals",
            "total_stock_on_hand_d, pct_SKU_stockout_d, avg_days_of_supply_d, n_overstock_d",
            "inventory end-of-month → reindex daily forward-fill. Test 2023+ → forward-fill 2022-12-31 snapshot hoặc drop. Decide cùng web_traffic.",
            "ý #5 Inventory-sales mismatch",
        ],
        [
            "Cash-flow features",
            "money_in = Σ credit_card + apple_pay + bank_transfer trên order_date + Σ cod trên delivery_date (khi tiền thực nhận)  ·  money_out = Σ refund_amount từ returns  ·  net_cash_d = money_in − money_out",
            "payments + shipments (delivery_date cho COD) + returns. Note: COD revenue được 'ghi' khi giao thành công, không phải khi đặt → proxy doanh thu thực nhận khác với Revenue đăng ký.",
            "Bổ trợ target Revenue/COGS + ý #7 Payment behaviour",
        ],
        [
            "Customer / Return signals (lag)",
            "lag_365_return_rate, lag_365_new_customer_share, new_customers_count (mỗi ngày số customer có first order = day t, chỉ dùng làm lag 365), average_ratings_d (rolling 30 ngày)",
            "returns + customers + reviews, chỉ dùng làm LAG 365 để tránh future leak. new_customers_count / average_ratings phải shift ≥548 ngày hoặc chỉ dùng historic pattern cho test period.",
            "ý #1 Lifecycle + ý #3 Return reason + ý #8 Review→return",
        ],
        [
            "Calendar features (SELF-BUILT)",
            "day_of_week, day_of_month, month, quarter, is_weekend, is_month_end, is_tet (lunar), is_quoc_khanh, is_black_friday, is_1111, is_1212",
            "⚠ KHÔNG import Prophet holidays (external data → LOẠI). Tự build DataFrame VN holiday 2012-2024 từ list tĩnh. day_of_month/is_weekend giúp bắt pattern pay-day + cuối tuần.",
            "ý #10 Seasonality (VN event list)",
        ],
        [
            "Model chính",
            "LightGBM 2 head (Revenue + COGS) — direct multi-output; Python 3.10, lightgbm ≥4.0, random_state=42",
            "Cân nhắc 2 model độc lập vs 1 MultiOutput. Early stopping theo temporal val. n_estimators cap 2000. MAE là metric chính (Kaggle), KHÔNG dùng MSE nội bộ.",
            "— (mô hình trung tâm)",
        ],
        [
            "Model candidates (taxonomy)",
            "Statistical: SARIMAX (Prophet ⚠ §7 rủi ro external holiday → skip hoặc tắt holidays)  ·  ML: LightGBM, XGBoost, CatBoost (direct multi-output)  ·  DL: DLinear, N-HITS, TimeXer, TimeMixer, TiDE, Temporal Fusion Transformer (note thời gian train + GPU cost)  ·  Zero-shot / Foundation: TimesFM, Moirai, Chronos",
            "Foundation weights KHÔNG phải 'external data' — chỉ dùng train-free inference trên sales.csv nội bộ, hợp lệ ràng buộc §3. DL options phải cân nhắc train time vs pred_len=548 (TFT/N-HITS thường >1h GPU).",
            "— (shortlist cho Phase 6c)",
        ],
        [
            "Ensemble strategies",
            "(a) Top 3-5 best models → weighted average với learnable weights (lr-grid trên val 2022)  ·  (b) Temporal hierarchical: train tại Daily / Weekly / Monthly / Quarterly / Yearly → chọn best per level → Reconciliation (Top-down / Bottom-up / MinT — tối thiểu implement Bottom-up vì đơn giản nhất)",
            "Robust hơn khi COVID-like shock. Weight chọn trên val set 2022. Hierarchical reconciliation giúp aggregate consistency giữa Daily và Monthly forecast.",
            "ý #10 Seasonality (SARIMAX bắt mùa vụ) + §4.1 nút (4)",
        ],
        [
            "Temporal CV — 5-fold rolling window",
            "k=365 step, pred_len=548  ·  weights = [0.1, 0.15, 0.2, 0.25, 0.3] (fold mới nhất weight cao hơn)  ·  Fold 1: Train [t_0 → t_n], Val [t_n+1 → t_n+pred_len]  ·  Fold 2: Train [t_0 → t_n+k], Val [t_n+k+1 → t_n+k+pred_len]  ·  …",
            "Tổng train 3,833 ngày → có thể chỉ đủ 2-3 fold thực tế với pred_len=548. Nếu thiếu, degrade xuống 3-fold và note trong báo cáo. KHÔNG random split. Code: src/cv.py (Phase 6).",
            "— (audit §3.2 bẫy #2 + §7 risk leakage)",
        ],
        [
            "Metric selection",
            "Primary = MAE (Kaggle leaderboard)  ·  Sanity-check = RMSE, R²  ·  Submit Kaggle bằng model có MAE thấp nhất trên val trung bình (weighted theo fold recency)",
            "KHÔNG dùng MSE nội bộ — MAE là metric Kaggle chính. Weighted MAE ưu tiên fold mới nhất để phản ánh pattern gần test period.",
            "— (căn cứ §2.3 Part 3 rubric hiệu suất)",
        ],
        [
            "Explainability",
            "SHAP TreeExplainer: summary plot + top-10 feature bar + 2 force plot (ngày bất thường)",
            "Diễn giải business: 'Tết → lag_365 dominate vì pattern năm trước'. Permutation importance bổ sung để cross-check SHAP.",
            "Narrative §4.1 nút (4) Dự báo + §4.4 (justify action bằng feature importance)",
        ],
    ],
)
add_para(
    doc,
    "→ Output của §4.3 quay ngược phục vụ §4.4: con số 'Dự kiến impact định lượng' trong §4.4 nên dùng forecast từ model này "
    "(vd: tính Δrevenue khi +5% price Premium = model predict với feature price tăng). "
    "Nếu thiếu hụt thời gian, dùng công thức heuristic trong §4.4 cột 'CSV & công thức'.",
    italic=True,
    size=10,
)

# ---------- 4.4 Prescriptive deep dive ----------
add_heading(doc, "4.4. Tầng Prescriptive — KPI & số định lượng mẫu", level=2)
add_para(doc, "→ Phụ trách: Kiên (Insights Lead). Số 'Dự kiến impact' nên dùng output từ §4.3 (model). Mỗi hàng là ending của 1 ý §4.2.", italic=True, size=10)
add_para(
    doc,
    "Top team được đánh giá ở đây (Insight 15đ). Cột 'CSV & công thức' liệt kê nguồn + metric cụ thể "
    "để phản biện khi BGK hỏi 'dựa vào đâu?'. Cột 'Dự kiến impact' cho con số ước lượng (phải định lượng được). "
    "Cột '← §4.2' chỉ rõ ý EDA nào sinh ra hàng này; cột 'Business goal' chỉ rõ 1 trong 3 mục tiêu §2.1 mà khuyến nghị này phục vụ.",
)

add_table(
    doc,
    headers=["Domain", "Khuyến nghị cụ thể", "CSV & công thức", "Dự kiến impact định lượng", "← §4.2 ý", "Business goal (§2.1)"],
    rows=[
        [
            "Pricing & discount",
            "Tăng giá segment Premium +5%; segment có demand elasticity < 0.5 trong analysis promo",
            "products (price, cogs, segment) + order_items + promotions → margin % + price-elasticity = %Δqty/%Δprice",
            "Nếu elasticity 0.3, +5% price → −1.5% qty → net +3.4% revenue Premium",
            "#11 Gross-margin × segment",
            "(2) Promo planning",
        ],
        [
            "Inventory reorder",
            "Reorder threshold = days_of_supply × 1.3 cho top 50 SKU (80% revenue)",
            "inventory + demand forecast từ sales → expected_daily_demand × 1.3",
            "Giảm stockout 30% → reclaim ~$X revenue/năm (tính từ lost_revenue analysis ý #5)",
            "#5 Inventory-sales mismatch + #10 Seasonality",
            "(1) Inventory optimization",
        ],
        [
            "Channel-mix budget",
            "Shift 20% budget social_media → paid_search (conv proxy cao hơn 2.3×)",
            "web_traffic + orders (order_source) → (n_orders/sessions) × AOV",
            "Ở cùng budget, +18% total orders từ paid_search (nếu ROAS scale tuyến tính)",
            "#6 Web-traffic funnel",
            "(2) Promo planning",
        ],
        [
            "Customer retention",
            "Email re-engagement sau ~90 ngày (median inter-order gap từ Q1 MCQ)",
            "customers + orders → cohort retention + inter-order gap distribution",
            "Re-engage 20% dormant cohort → $Y incremental nếu AOV = $Z",
            "#1 Lifecycle + #8 Review→return",
            "(2) Promo planning",
        ],
        [
            "Regional logistics",
            "Ưu tiên warehouse/3PL ở rising region (East); min_order_value cho free-ship để filter casual",
            "orders + geography + shipments → revenue per region YoY + avg lead_time (delivery_date − ship_date) per region",
            "Giảm lead-time East từ 5 ngày → 3 ngày → tăng conversion trang checkout ~+4%",
            "#4 Regional growth + #12 Shipping-fee × return",
            "(3) Logistics",
        ],
        [
            "Size-guide fix",
            "Update size chart cho top 3 SKU có wrong_size return cao nhất; CSKH reach-out trong 48h với review ≤2",
            "returns (reason='wrong_size', refund_amount) + products (size, category) + reviews (rating)",
            "Giảm 30% wrong_size → refund_saved = N_wrong × avg_refund × 0.3",
            "#3 Return reason × size + #8 Review→return",
            "(1) Inventory optimization",
        ],
        [
            "Promo scheduling",
            "KHÔNG chạy promo trong tuần Tết +7 ngày (demand đã peak); target promo vào segment chưa mua 90 ngày",
            "promotions + sales + VN calendar self-built → uplift % promo vs non-promo trong/ngoài mùa",
            "Tiết kiệm discount cost 100% tuần Tết khi revenue vẫn peak tự nhiên",
            "#2 Promo ROI + #9 Cannibalization + #10 Seasonality",
            "(2) Promo planning",
        ],
        [
            "COD optimization",
            "Yêu cầu voucher trả trước 5% cho đơn COD trên $M value",
            "orders (payment_method, order_status) + payments (payment_value)",
            "Giảm cancel_rate COD từ 30% → 25% → reclaim $W revenue/năm",
            "#7 Payment/device behaviour",
            "(2) Promo planning",
        ],
    ],
)
add_para(
    doc,
    "→ Phản chiếu: mỗi hàng ở đây PHẢI có 1 chart EDA tương ứng trong §4.2 + 1 feature tương ứng trong §4.3 "
    "(nếu relevant). Nếu hàng nào thiếu chart EDA → rơi điểm 'Insight kinh doanh' (thiếu chứng cứ); nếu thiếu feature model → "
    "rơi điểm 'Explainability' trong Phần 3. Dùng audit §4.5 để kiểm tra cross-coverage.",
    italic=True,
    size=10,
)

# ---------- 4.5 Coverage matrix ----------
add_heading(doc, "4.5. Dataset coverage matrix", level=2)
add_para(doc, "→ Audit cuối cùng: cột 'Ý §4.2' chỉ rõ ý nào sẽ 'chạm' đến CSV đó → đảm bảo không CSV nào bị bỏ quên.", italic=True, size=10)
add_para(
    doc,
    "Ma trận 14 CSV × 4 tầng phân tích — giúp team chắc chắn đã 'đụng' đủ dataset trong bài nộp. "
    "Ký hiệu: ✓ = bắt buộc phủ; — = không cần (file chỉ dùng cho Analytical target).",
)
add_table(
    doc,
    headers=["CSV", "D (What)", "Di (Why)", "P (Likely)", "Pr (Action)", "Ý §4.2 phủ"],
    rows=[
        ["products", "✓", "✓", "—", "✓", "#3, #11"],
        ["customers", "✓", "✓", "✓", "✓", "#1, #9"],
        ["geography", "✓", "✓", "—", "✓", "#4"],
        ["promotions", "—", "✓", "✓", "✓", "#2, #9, #10"],
        ["orders", "✓", "✓", "✓", "✓", "#1, #4, #6, #7, #9, #12"],
        ["order_items", "✓", "✓", "✓", "✓", "#2, #3, #4, #5, #9, #11"],
        ["payments", "✓", "✓", "—", "✓", "#7"],
        ["shipments", "✓", "✓", "—", "✓", "#4, #12"],
        ["returns", "✓", "✓", "—", "✓", "#3, #8, #12"],
        ["reviews", "✓", "✓", "—", "—", "#8"],
        ["inventory", "✓", "✓", "✓", "✓", "#5"],
        ["web_traffic", "✓", "✓", "✓", "✓", "#6"],
        ["sales", "✓", "—", "✓", "✓", "#10 + nền §4.3"],
        ["sample_submission", "—", "—", "✓", "—", "— (target format Phần 3)"],
    ],
)
add_para(
    doc,
    "→ Nếu cột 'Ý §4.2 phủ' của CSV nào rỗng khi đọc bảng này → quay lại §4.2 thêm ý mới hoặc mở rộng ý hiện có. "
    "Hiện tại 14/14 CSV có ít nhất 1 ý EDA — ổn để nộp.",
    italic=True,
    size=10,
)

# ---------- 5. Plan 12 ngày ----------
add_heading(doc, "5. Kế hoạch 12 ngày → deadline 2026-05-01", level=1)
add_table(
    doc,
    headers=["Giai đoạn", "Ngày", "Output", "Lead"],
    rows=[
        ["Re-audit + Data audit", "04-19 → 04-20", "Re-audit PDF ✅; 01_data_audit.ipynb (schema, missing, outliers, integrity)", "All (Đồng dẫn)"],
        ["MCQ", "04-21", "02_mcq.ipynb + outputs/mcq_answers.json (verify chéo)", "Hiển"],
        ["EDA Descriptive + Diagnostic", "04-21 → 04-23", "03_eda_descriptive.ipynb — tầng 1-2, 6-8 chart đắt giá (§4.2 ý #1/#3/#4/#5/#6/#7/#11)", "Đồng (Lead Analyst)"],
        ["EDA Predictive + Prescriptive", "04-24 → 04-26", "04_eda_predictive.ipynb — tầng 3-4, story-led (§4.1) + action định lượng (§4.4)", "Kiên (Insights Lead)"],
        ["Feature engineering + Model", "04-27 → 04-29", "05_FE.ipynb, 06_baseline.ipynb, 07_advanced.ipynb (§4.3 feature blocks) + Kaggle submission đầu tiên", "Phúc (ML Engineer)"],
        ["Kaggle iteration + Report", "04-30", "2-3 Kaggle submission refinement + main.tex viết song song (Hiển lay-out, Kiên viết EDA section, Phúc viết Model section)", "Phúc + Hiển (lead), Kiên hỗ trợ"],
        ["Finalize + GitHub + form", "05-01", "main.pdf ≤4 trang (NeurIPS), README.md, submit form chính thức (§8)", "Hiển + All"],
    ],
)

add_callout(
    doc,
    "Dự phòng: Buffer 1 ngày (05-01 sáng) cho việc submit cuối + fix lỗi format PDF/Kaggle. "
    "Không lùi deadline nội bộ sang 04-30 trừ khi có rủi ro lớn. Submission Kaggle phải đúng 548 dòng, "
    "thứ tự khớp 100% sample_submission.csv (Kaggle sẽ reject nếu sai format).",
)

add_callout(
    doc,
    "Rebuild baseline: baseline.ipynb hiện dùng DATA_DIR='dataset/' — sai path so với folder dataset-datathon-2026-round-1/. "
    "Trong giai đoạn 04-27 → 04-29 (ML Engineer), bước đầu tiên là fork baseline sang 06_model_baseline.ipynb, sửa DATA_DIR, "
    "re-run end-to-end để xác nhận reproducibility trước khi build LightGBM.",
)

# ---------- 6. Roles ----------
add_heading(doc, "6. Phân công vai trò (đã chốt 2026-04-19)", level=1)
add_para(
    doc,
    "Team đã chốt 4 role. Các section khác trong file này tham chiếu bằng tên role (Lead Analyst / Insights Lead / ML Engineer / Đội trưởng) — xem cột '→ Liên quan' để biết section nào của file phục vụ cho role nào.",
)
add_table(
    doc,
    headers=["Role", "Người", "Trách nhiệm chính", "Notebook", "→ Liên quan trong file"],
    rows=[
        [
            "Đội trưởng",
            "Hiển",
            "Coordinator, MCQ, LaTeX báo cáo, GitHub, submission form, theo deadline.",
            "02, report/, README",
            "§2.2 (MCQ), §5 (04-21 + 04-30 + 05-01), §8 (checklist nộp bài)",
        ],
        [
            "Lead Analyst",
            "Đồng",
            "EDA tầng 1-2 (Descriptive + Diagnostic), data audit, storytelling nền tảng.",
            "01_data_audit, 03_eda_descriptive",
            "§2.3 (tầng D+Di), §3.1–§3.2 (cột đắt + bẫy dữ liệu), §4.2 ý #1/#3/#4/#5/#6/#7/#11 (tag D/Di), §4.5 coverage, §5 (04-19→04-23)",
        ],
        [
            "Insights Lead",
            "Kiên",
            "EDA tầng 3-4 (Predictive + Prescriptive), viết insight + storytelling cho báo cáo.",
            "04_eda_predictive",
            "§2.3 (tầng P+Pr), §4.1 narrative, §4.2 ý #2/#8/#9/#10/#12 (tag P/Pr), §4.4 toàn bộ (KPI + impact), §5 (04-24→04-26)",
        ],
        [
            "ML Engineer",
            "Phúc",
            "Feature engineering, baseline, LightGBM, SHAP, Kaggle submissions.",
            "05_feature_engineering, 06_model_baseline, 07_model_advanced",
            "§2.4 (ràng buộc Phần 3), §3.2 bẫy #2/#3/#5/#7 (leakage), §4.3 toàn bộ feature blocks, §4.4 cột 'Dự kiến impact' (input cho forecasting), §5 (04-27→04-30), §7 rủi ro leakage/Prophet/RAM",
        ],
    ],
)
add_callout(
    doc,
    "Luật review chéo (bắt buộc): (a) Mỗi notebook có cell đầu tiên set SEED=42 + in phiên bản pandas/numpy/lightgbm để debug cross-machine. "
    "(b) Hiển review MCQ+Report; Đồng review EDA Predictive của Kiên; Kiên review Prescriptive-narrative trong model report của Phúc; Phúc review feature Engineering pipeline của mình + audit leakage trên notebook 03/04 trước khi đưa vào model. "
    "(c) Trước mỗi deadline nội bộ (§5), 1 thành viên chạy end-to-end trên máy thứ 2 để xác nhận reproducibility.",
)

# ---------- 7. Risks ----------
add_heading(doc, "7. Rủi ro & mitigation", level=1)
add_table(
    doc,
    headers=["Rủi ro", "Tác động", "Mitigation"],
    rows=[
        ["RAM laptop không đủ cho orders.csv (44MB) + order_items.csv (23MB)", "Notebook crash khi load full", "Khai báo dtype (int32 thay int64, category cho text), dùng chunksize, hoặc move sang Kaggle/Colab."],
        ["Prophet auto-load VN holiday → bị coi 'external data'", "Loại Phần 3 (mất 20đ)", "KHÔNG dùng Prophet với holidays=... Tự xây VN calendar tĩnh 2012-2024."],
        ["Dùng Revenue/COGS 2023+ làm feature (leakage)", "Loại Phần 3", "Anti-leakage checklist CLAUDE.md §9. Code review cross check. Unit test: X_train.max(Date) < '2023-01-01'."],
        ["web_traffic / inventory không có cho 2023+", "Mất feature mạnh", "3 option: forecast riêng / rolling proxy / drop. Decide Phase 6 đồng thời cho cả 2 file."],
        ["PDF báo cáo ≤4 trang RẤT chặt cho cả EDA + Model", "Cắt insight", "Plan layout từ tuần 1; ưu tiên 5-6 chart đắt giá nhất + bảng tóm tắt; appendix cho phụ."],
        ["Repository public → ảnh thẻ SV bị leak nếu commit nhầm", "Vi phạm bảo mật cá nhân", "Thêm photos/, .env vào .gitignore. KHÔNG push ảnh thẻ. Upload ảnh riêng qua form."],
        ["Screen recording .mov 5GB", "Push fail, repo bloat", "Thêm *.mov vào .gitignore."],
        ["Kaggle team chưa merge trước deadline", "Submission không tính cho team", "Hiển check ngay 04-19: join competition → team → merge 4 thành viên."],
        ["Submission sai format (thiếu cột, sai order)", "Kaggle reject, mất 0 điểm", "Assert 548 dòng; assert columns == ['Date','Revenue','COGS']; assert Date order match sample_submission."],
    ],
)

# ---------- 8. Checklist nộp bài ----------
add_heading(doc, "8. Checklist nộp bài", level=1)
add_para(doc, "Tick từng mục trước khi bấm Submit form (theo PDF page 16):")
checklist = [
    "Kaggle submission đúng 548 dòng, thứ tự khớp 100% sample_submission.csv, 3 cột Date/Revenue/COGS.",
    "GitHub repo public hoặc invite organizers; README.md có hướng dẫn reproduce đầy đủ + requirements.txt + SEED=42.",
    "Báo cáo PDF ≤4 trang (không tính refs/appendix), template NeurIPS 2025, có link GitHub trong nội dung PDF.",
    "Báo cáo có mục Explainability (SHAP / feature importance / PDP) bằng ngôn ngữ business.",
    "Báo cáo có section riêng cho EDA (Phần 2) và Model (Phần 3).",
    "Form: đáp án 10 MCQ + upload PDF + GitHub link + Kaggle submission link.",
    "Form: ảnh thẻ sinh viên của TẤT CẢ 4 thành viên.",
    "Form: tickbox cam kết ≥1 thành viên tham dự Vòng Chung kết 2026-05-23 tại VinUni HN.",
    "Random seed = 42 set consistent ở mọi notebook (np, random, lightgbm random_state).",
    "KHÔNG có Revenue/COGS test set xuất hiện trong feature pipeline (verify bằng assert).",
    "KHÔNG dùng dữ liệu/library auto-load external data (Prophet holiday, Yahoo Finance, etc.).",
    "Test reproduce end-to-end trên máy thứ 2 ít nhất 1 lần trước deadline.",
]
for item in checklist:
    add_bullet(doc, item)

# ---------- 9. Open questions ----------
add_heading(doc, "9. Câu hỏi mở cho cuộc họp đầu tiên", level=1)
add_callout(
    doc,
    "✅ Đã chốt (2026-04-19) + re-assign role Phúc↔Kiên (2026-04-21) — chi tiết §6. Hiện tại: Hiển (Đội trưởng) · Đồng (Lead Analyst) · Kiên (Insights Lead — EDA P+Pr) · Phúc (ML Engineer — Model).",
    bg="E6F4EA",
)
add_para(doc, "Còn lại — cần thống nhất trong cuộc họp team đầu tiên (04-19 / 04-20):")
questions = [
    ("GitHub repo", "Tên repo? Owner account? Public từ đầu hay private rồi public sát deadline?"),
    ("Kaggle team setup", "Tất cả 4 thành viên đã join competition chưa? Merge team trên Kaggle hay 1 account submit?"),
    ("Môi trường compute", "Ai dùng laptop? RAM bao nhiêu? Có cần Colab/Kaggle Notebook làm fallback không?"),
    ("Web traffic test period", "Chọn option (a) self-forecast / (b) rolling mean proxy / (c) bỏ feature?"),
    ("Inventory test period", "Snapshot dừng 2022-12-31, test 18 tháng không có. Forward-fill 2022-12 hay drop? (Nên quyết cùng lúc với web_traffic.)"),
    ("Prophet vs LightGBM", "Có dùng Prophet không (rủi ro external data)? Hay LightGBM-only ngay từ đầu?"),
    ("Lịch họp", "Họp chốt mỗi mấy ngày (đề xuất: 04-20, 04-23, 04-27, 04-30)?"),
    ("Template chart", "Chọn matplotlib + seaborn hay plotly? Phải đồng nhất style (color palette, title size) — đề xuất seaborn whitegrid + colorblind palette."),
    ("Evaluation budget Kaggle", "Submission limit/ngày? Lên lịch submit ưu tiên model mạnh nhất, tránh burn limit vào sanity check."),
]
for title, body in questions:
    p = doc.add_paragraph(style="List Bullet")
    r1 = p.add_run(f"{title} — ")
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(body)
    r2.font.size = Pt(11)

# ---------- Appendix ----------
add_heading(doc, "Phụ lục — File tham chiếu nội bộ", level=1)
add_table(
    doc,
    headers=["File", "Vai trò"],
    rows=[
        ["CLAUDE.md", "Brain file tổng cho Claude sessions — đọc §1–§13 trước khi làm việc. Luôn là source of truth cho dự án."],
        ["Đề thi Vòng 1.pdf (2026-04-19)", "Đề chính thức — luôn là nguồn cuối cùng nếu có tranh cãi. Bản cập nhật mới nhất."],
        ["Rule_book.docx", "Quy chế chi tiết, ít chi tiết hơn PDF nhưng có business context."],
        ["dataset-datathon-2026-round-1/baseline.ipynb", "Baseline seasonal × growth — dùng làm benchmark tối thiểu."],
        ["dataset-datathon-2026-round-1/dataset_description.docx", "Data dictionary chi tiết — cross-check với PDF §1."],
        ["scripts/build_team_brief_v2.py", "Script sinh lại file docx này (reproducible)."],
    ],
)

add_para(
    doc,
    "Tài liệu này là snapshot v2 tại 2026-04-19. Sau cuộc họp đầu, update lại §6 (phân công), §9 (resolved questions), "
    "và lưu thành phiên bản v3 nếu cần. Lịch sử thay đổi lớn giữa v1 và v2:",
    italic=True,
)
add_bullet(doc, "Sync sang 14 file CSV (PDF Table 1) thay vì 15; xoá mention inventory_enhanced.csv obsolete.")
add_bullet(doc, "Thêm chi tiết methodology cho từng MCQ (phương pháp tính pseudocode).")
add_bullet(doc, "Mở rộng mỗi ý tưởng EDA thành (H) Hypothesis / (T) Test / (A) Action với con số định lượng.")
add_bullet(doc, "Chi tiết hoá feature block model: lag values cụ thể, expanding-window CV folds, SHAP plan.")
add_bullet(doc, "Thêm cột 'Dự kiến impact định lượng' cho toàn bộ Prescriptive recommendations.")
add_bullet(doc, "Thêm bẫy dữ liệu #7 (inventory monthly snapshot) và nhắc PDF dùng 2 tên cho file train.")
add_bullet(doc, "Timeline đổi sang 12 ngày (từ 04-19) với buffer submit cuối.")

doc.save(OUT)
print(f"Saved: {OUT}")
print(f"Size: {OUT.stat().st_size / 1024:.1f} KB")

"""
Generate the team-brief .docx for Datathon 2026 Round 1.

Output: ../Vòng 1 — Phân tích đề & Brainstorm (team brief).docx
Run:    python3 scripts/gen_brief.py
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

OUTPUT = Path(__file__).resolve().parent.parent / "Vòng 1 — Phân tích đề & Brainstorm (team brief).docx"

NAVY = RGBColor(0x0B, 0x2E, 0x5C)
ACCENT = RGBColor(0xC0, 0x39, 0x2B)
GREY = RGBColor(0x55, 0x55, 0x55)


def set_cell_shading(cell, hex_fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tc_pr.append(shd)


def style_run(run, *, size=11, bold=False, color=None, italic=False):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color


def add_para(doc, text, *, size=11, bold=False, color=None, italic=False, align=None, space_after=4):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    style_run(p.add_run(text), size=size, bold=bold, color=color, italic=italic)
    return p


def add_heading(doc, text, level):
    sizes = {1: 18, 2: 14, 3: 12}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level > 1 else 14)
    p.paragraph_format.space_after = Pt(4)
    style_run(p.add_run(text), size=sizes[level], bold=True, color=NAVY)
    if level == 1:
        bdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), "8")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), "0B2E5C")
        bdr.append(bot)
        p._p.get_or_add_pPr().append(bdr)


def add_bullets(doc, items, *, size=11, indent=0):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.5 + indent)
        p.paragraph_format.space_after = Pt(2)
        if isinstance(it, tuple):
            head, body = it
            r1 = p.add_run(head)
            style_run(r1, size=size, bold=True)
            r2 = p.add_run(" — " + body)
            style_run(r2, size=size)
        else:
            style_run(p.add_run(it), size=size)


def add_table(doc, header, rows, *, widths=None, header_fill="0B2E5C", header_white=True, font_size=10):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = "Light Grid Accent 1"
    t.autofit = False
    if widths:
        for i, w in enumerate(widths):
            for cell in t.columns[i].cells:
                cell.width = Cm(w)
    for i, h in enumerate(header):
        cell = t.cell(0, i)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, header_fill)
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        style_run(run, size=font_size, bold=True,
                  color=RGBColor(0xFF, 0xFF, 0xFF) if header_white else NAVY)
    for r, row in enumerate(rows, start=1):
        for c, val in enumerate(row):
            cell = t.cell(r, c)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            style_run(run, size=font_size)
    return t


def add_callout(doc, title, body, color=ACCENT):
    t = doc.add_table(rows=1, cols=1)
    t.autofit = False
    cell = t.cell(0, 0)
    cell.width = Cm(17)
    set_cell_shading(cell, "FFF4E5")
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    style_run(p.add_run(title + ": "), size=11, bold=True, color=color)
    style_run(p.add_run(body), size=11)
    doc.add_paragraph()


def build():
    doc = Document()

    # Page setup A4
    for section in doc.sections:
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ========== COVER ==========
    add_para(doc, "DATATHON 2026 — VÒNG 1", size=14, bold=True, color=NAVY,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, "The Gridbreakers · Breaking Business Boundaries", size=11, italic=True,
             color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
    add_para(doc, "PHÂN TÍCH ĐỀ BÀI & BRAINSTORM CHIẾN LƯỢC", size=22, bold=True, color=NAVY,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    add_para(doc, "Tài liệu kickoff team — chia sẻ trước cuộc họp đầu tiên",
             size=12, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)

    add_table(
        doc,
        ["Mục", "Thông tin"],
        [
            ["Đội", "The Gridbreakers (4 thành viên: Hiển — đội trưởng, Phúc, Kiên, Đồng)"],
            ["Ngày phát hành brief", "2026-04-18"],
            ["Deadline nộp Vòng 1", "2026-05-01 (cứng — còn ~13 ngày)"],
            ["Vòng Chung kết", "2026-05-23 tại VinUniversity, Hà Nội"],
            ["Kaggle", "https://www.kaggle.com/competitions/datathon-2026-round-1"],
            ["LaTeX template", "NeurIPS 2025 — https://neurips.cc/Conferences/2025/CallForPapers"],
            ["Vai trò giả định", "Data Scientist tại công ty thời trang TMĐT Việt Nam"],
            ["Dataset", "15 file CSV (~130MB), 10.5 năm (2012-07 → 2022-12)"],
        ],
        widths=[5, 12],
    )

    doc.add_page_break()

    # ========== 1. TL;DR ==========
    add_heading(doc, "1. Tóm tắt nhanh (TL;DR)", 1)
    add_para(doc,
             "Datathon 2026 Vòng 1 yêu cầu chúng ta đóng vai Data Scientist của 1 công ty thời trang TMĐT VN, "
             "khai thác 15 file CSV (10.5 năm dữ liệu) để biến dữ liệu thành giải pháp kinh doanh. "
             "Bài thi gồm 3 phần với tổng 100 điểm; trọng số áp đảo nằm ở Phần 2 (EDA = 60đ).")
    add_table(
        doc,
        ["Phần", "Nội dung", "Điểm", "Tỷ trọng", "Ưu tiên đầu tư"],
        [
            ["1", "MCQ — 10 câu × 2đ", "20", "20%", "Trung bình (1 ngày)"],
            ["2", "EDA & Data Visualization", "60", "60%", "★★★ CAO NHẤT (5–6 ngày)"],
            ["3", "Sales Forecasting (Kaggle)", "20", "20%", "★★ Cao (3–4 ngày)"],
            ["", "Tổng", "100", "100%", ""],
        ],
        widths=[1.2, 7.5, 1.6, 1.8, 4.5],
    )
    add_callout(
        doc, "Insight chiến lược",
        "Phần 2 EDA đáng giá gấp 3 lần Phần 3 Forecasting. Kaggle ranking chỉ chiếm 12/100 điểm "
        "tổng — không cần 'all-in' vào model. Đầu tư mạnh nhất vào EDA storytelling + 4 tầng "
        "phân tích (đặc biệt Prescriptive). Phần 3 cần beat baseline & viết explainability tốt."
    )

    # ========== 2. PHÂN TÍCH ĐỀ ==========
    add_heading(doc, "2. Phân tích đề bài", 1)

    # 2.1 Bối cảnh
    add_heading(doc, "2.1. Bối cảnh kinh doanh", 2)
    add_para(doc,
             "Doanh nghiệp cần dự báo doanh thu chính xác ở mức chi tiết theo ngày để đạt 3 mục tiêu cốt lõi:")
    add_bullets(doc, [
        ("Tối ưu hoá phân bổ tồn kho",
         "biết bán bao nhiêu để nhập đủ, không thừa — tránh stockout & overstock."),
        ("Lập kế hoạch khuyến mãi",
         "chọn đúng thời điểm chạy promo theo dự báo demand."),
        ("Quản lý logistics toàn quốc",
         "phân bổ hàng giữa các vùng (East / Central / West)."),
    ])
    add_para(doc,
             "→ Mọi insight EDA và feature model phải kết nối ngược về 1 trong 3 mục tiêu này khi viết phần "
             "Prescriptive. Đây là tiêu chí ban giám khảo đánh giá 'business value'.",
             italic=True, color=GREY)

    # 2.2 Phần 1
    add_heading(doc, "2.2. Phần 1 — MCQ (20 điểm)", 2)
    add_bullets(doc, [
        "10 câu × 2đ = 20đ. KHÔNG trừ điểm câu sai → luôn đoán cả khi không chắc.",
        "Mỗi câu chỉ chọn 1 đáp án.",
        "Tất cả câu hỏi đều tính trực tiếp từ CSV → phải code thật để verify.",
    ])
    add_para(doc, "10 câu hỏi & file nguồn cần truy vấn:", bold=True)
    add_table(
        doc,
        ["Q", "Đề ngắn", "File chính cần join", "Lưu ý"],
        [
            ["1", "Trung vị inter-order gap (khách >1 đơn)", "orders.csv", "Sort theo customer_id, order_date; diff."],
            ["2", "Segment có gross margin cao nhất", "products.csv", "Margin = (price − cogs)/price."],
            ["3", "Lý do trả hàng phổ biến nhất ở Streetwear", "returns + products", "Join theo product_id, filter category."],
            ["4", "traffic_source có bounce_rate trung bình thấp nhất", "web_traffic.csv", "Trung bình theo từng source."],
            ["5", "% dòng order_items có promo_id ≠ null", "order_items.csv", "1 phép tính count/total."],
            ["6", "age_group có số đơn TB / khách cao nhất", "customers + orders", "Tổng orders / unique customers."],
            ["7", "Region có tổng doanh thu cao nhất", "order_items + orders + geography", "Revenue = Σ(qty×unit_price) từ order_items; join orders.zip → geography.region. Đề gọi sales_train.csv nhưng file thực tế là sales.csv."],
            ["8", "Payment method phổ biến nhất ở đơn cancelled", "orders.csv", "Filter status = 'cancelled'."],
            ["9", "Size có return rate cao nhất", "returns + order_items + products", "rate = #returns / #order_items lines."],
            ["10", "Installment plan có giá trị TT TB cao nhất", "payments.csv", "Group by installments, mean(payment_value)."],
        ],
        widths=[0.7, 5.3, 4.0, 6.5],
        font_size=9,
    )
    add_callout(
        doc, "Chiến thuật MCQ",
        "1 người (đề xuất Hiển) viết 1 notebook 02_mcq.ipynb chạy cả 10 câu chỉ trong vài giờ. "
        "Save kết quả vào file riêng + share screenshot cho team verify chéo."
    )

    # 2.3 Phần 2
    add_heading(doc, "2.3. Phần 2 — EDA & Data Visualization (60 điểm)", 2)
    add_para(doc,
             "Đây là phần ĂN ĐIỂM NHẤT. Không có đáp án đúng duy nhất — ban giám khảo chấm khả năng "
             "kể chuyện bằng dữ liệu (data storytelling) và chiều sâu phân tích. Bài nộp gồm 2 phần: "
             "(1) Visualizations có title/axis/legend rõ; (2) Phần Analysis đi kèm mỗi chart, gồm "
             "key findings + business implications + actionable recommendations.")
    add_para(doc, "Rubric chính thức (PDF tr.13-14):", bold=True)
    add_table(
        doc,
        ["Tiêu chí", "Mô tả", "Điểm tối đa"],
        [
            ["Chất lượng trực quan hoá", "Chart có title/axis/legend đủ, chọn loại chart phù hợp, thẩm mỹ", "15"],
            ["Chiều sâu phân tích", "Phủ đủ 4 tầng Descriptive → Diagnostic → Predictive → Prescriptive", "25"],
            ["Insight kinh doanh", "Đề xuất cụ thể, định lượng, áp dụng được ngay", "15"],
            ["Tính sáng tạo & kể chuyện", "Góc nhìn độc đáo, narrative xuyên suốt, kết nối nhiều bảng", "5"],
            ["", "Tổng", "60"],
        ],
        widths=[4.5, 10.5, 2.0],
    )

    add_para(doc, "4 tầng phân tích bắt buộc — ví dụ câu hỏi cho mỗi tầng:", bold=True)
    add_table(
        doc,
        ["Tầng", "Câu hỏi BGK", "Ví dụ phân tích"],
        [
            ["Descriptive", "What happened?", "Doanh thu YoY 2013–2022, top categories, geo distribution."],
            ["Diagnostic", "Why did it happen?", "Mùa COVID 2020-2021 doanh thu giảm vì sao? So sánh region/segment."],
            ["Predictive", "What is likely?", "Trend & seasonality phân rã, dự báo tồn kho theo segment."],
            ["Prescriptive", "What should we do?", "Tăng budget paid_search ở East do conversion 2.3× (defensible)."],
        ],
        widths=[2.5, 4.0, 10.5],
    )
    add_callout(
        doc, "Quy tắc vàng",
        "Top team đạt PRESCRIPTIVE NHẤT QUÁN trên nhiều phân tích sẽ được điểm cao nhất. "
        "Mỗi visualization phải có ít nhất 1 dòng 'So what?' + 1 dòng 'Action gì?'."
    )

    # 2.4 Phần 3
    add_heading(doc, "2.4. Phần 3 — Sales Forecasting (20 điểm)", 2)
    add_bullets(doc, [
        ("Target", "Predict Revenue (& COGS) hàng ngày cho 548 ngày 2023-01-01 → 2024-07-01."),
        ("Train", "sales.csv 04/07/2012 → 31/12/2022 (3,833 ngày)."),
        ("Submission", "submission.csv 3 cột (Date, Revenue, COGS), thứ tự dòng KHỚP 100% sample_submission."),
        ("Metrics", "MAE, RMSE (càng thấp càng tốt), R² (càng cao càng tốt, ~1)."),
    ])
    add_para(doc, "Quy đổi điểm Phần 3 (PDF tr.15-16):", bold=True)
    add_table(
        doc,
        ["Thành phần", "Band", "Mô tả"],
        [
            ["Hiệu suất mô hình (12đ)", "10–12", "Top leaderboard; MAE/RMSE thấp, R² cao"],
            ["", "5–9", "Trung bình; mô hình hoạt động nhưng chưa tối ưu"],
            ["", "3–4", "Hợp lệ nhưng performance thấp (mức điểm sàn)"],
            ["Báo cáo kỹ thuật (8đ)", "7–8", "Pipeline rõ, temporal CV đúng, SHAP cụ thể, tuân thủ ràng buộc"],
            ["", "4–6", "Pipeline đủ dùng, giải thích định tính, vài ràng buộc chưa xử lý tường minh"],
            ["", "0–3", "Thiếu giải thích, không kiểm soát leakage, không tái lập"],
        ],
        widths=[5.0, 1.8, 10.2],
    )
    add_callout(
        doc, "Điều kiện LOẠI Phần 3 (mất nguyên 20đ)",
        "(1) Dùng Revenue/COGS test làm feature; (2) Dùng dữ liệu ngoài 15 CSV; "
        "(3) Không nộp code hoặc kết quả không tái lập."
    )

    doc.add_page_break()

    # ========== 3. DATASET TOUR ==========
    add_heading(doc, "3. Dataset tour — 15 CSV chia 4 lớp", 1)
    add_para(doc, "Đề chính thức (PDF §1.1) tổ chức 15 file thành 4 lớp logic. Đây là khung tham chiếu "
             "khi viết EDA — mỗi lớp có vai trò khác nhau:")
    add_table(
        doc,
        ["Lớp", "Vai trò", "File", "Đặc điểm"],
        [
            ["Master", "Tham chiếu, ít thay đổi",
             "products, customers, promotions, geography",
             "Slowly Changing Dimensions; dùng để dim & filter."],
            ["Transaction", "Sự kiện thô (event-level)",
             "orders, order_items, payments, shipments, returns, reviews",
             "Khối lượng lớn (~647K orders, 715K items). Gốc của mọi metric daily."],
            ["Analytical", "Aggregate đã pre-compute",
             "sales (train), sample_submission (test format)",
             "Daily Revenue + COGS — chính là TARGET."],
            ["Operational", "Vận hành nội bộ",
             "inventory (monthly snapshot), web_traffic (daily)",
             "Inventory chỉ end-of-month; web_traffic không có cho test period."],
        ],
        widths=[2.0, 3.5, 5.0, 6.5],
    )

    add_para(doc, "Quan hệ giữa các bảng (cardinality):", bold=True)
    add_table(
        doc,
        ["Quan hệ", "Cardinality"],
        [
            ["orders ↔ payments", "1 : 1"],
            ["orders ↔ shipments", "1 : 0 hoặc 1 (chỉ status shipped/delivered/returned)"],
            ["orders ↔ returns", "1 : 0 hoặc nhiều"],
            ["orders ↔ reviews", "1 : 0 hoặc nhiều (~20% delivered orders có review)"],
            ["order_items ↔ promotions", "nhiều : 0 hoặc 1 (qua promo_id, promo_id_2)"],
            ["products ↔ inventory", "1 : nhiều (1 dòng/sản phẩm/tháng)"],
        ],
        widths=[6.0, 11.0],
    )

    # §3bis — Quick-ref cột quan trọng
    add_heading(doc, "3.1. Quick-reference — cột đắt giá của mỗi CSV", 2)
    add_para(doc,
             "Dùng khi join 2+ bảng — khỏi phải mở dataset_description.docx. "
             "PK = primary key, FK = foreign key. Cột in đậm = thường làm feature hoặc groupby.")
    add_table(
        doc,
        ["File", "PK / FK", "Cột đắt giá cho phân tích"],
        [
            ["products", "PK product_id",
             "category (4), segment (8), size (S/M/L/XL), color, price, cogs"],
            ["customers", "PK customer_id, FK zip",
             "signup_date, age_group (nullable), gender (nullable), acquisition_channel (nullable)"],
            ["geography", "PK zip",
             "region (East/Central/West), city (42), district"],
            ["promotions", "PK promo_id",
             "promo_type (percentage/fixed), discount_value, start_date, end_date, applicable_category, stackable_flag, min_order_value"],
            ["orders", "PK order_id, FK customer_id, zip",
             "order_date, order_status (6), payment_method (5), device_type (3), order_source (6)"],
            ["order_items", "PK (order_id, product_id); FK promo_id, promo_id_2",
             "quantity, unit_price (ĐÃ giảm giá!), discount_amount"],
            ["payments", "FK order_id (1:1)",
             "payment_method, payment_value, installments (1/3/6/12)"],
            ["shipments", "FK order_id (1:0/1)",
             "ship_date, delivery_date, shipping_fee (0 = free)"],
            ["returns", "PK return_id; FK order_id, product_id",
             "return_date, return_reason (5), return_quantity, refund_amount"],
            ["reviews", "PK review_id; FK order_id, product_id, customer_id",
             "review_date, rating (1-5), review_title"],
            ["inventory", "PK (snapshot_date, product_id)",
             "stock_on_hand, stockout_days, days_of_supply, fill_rate, stockout_flag, overstock_flag, reorder_flag, sell_through_rate (+category, segment denorm)"],
            ["web_traffic", "PK date (2013-01-01+)",
             "sessions, unique_visitors, page_views, bounce_rate, avg_session_duration_sec, traffic_source (6). KHÔNG có conversion_rate → tự tính orders/sessions."],
            ["sales", "PK Date (2012-07-04 → 2022-12-31, 3833 dòng)",
             "Revenue, COGS — ★ TARGET train"],
            ["sample_submission", "PK Date (2023-01-01 → 2024-07-01, 548 dòng)",
             "Revenue, COGS dummy — ★ format test"],
        ],
        widths=[2.8, 4.5, 9.7],
        font_size=9,
    )

    add_callout(
        doc, "6 cái bẫy dữ liệu PHẢI biết",
        "(1) order_items.unit_price ĐÃ là giá sau giảm — không phải giá list; muốn ra giá gốc: "
        "(qty×unit_price + discount_amount) / qty. "
        "(2) Đề gọi sales_train.csv nhưng file thực tế tên sales.csv. "
        "(3) Đề liệt kê inventory_enhanced.csv nhưng file này KHÔNG tồn tại — inventory.csv hiện tại đã có cột enhanced. "
        "(4) web_traffic.csv chỉ có đến 2022-12-31 — test period 2023-01 → 2024-07 không có. "
        "(5) shipments không tồn tại cho orders status cancelled/paid/created. "
        "(6) web_traffic.csv bắt đầu 2013-01-01, trong khi sales.csv bắt đầu 2012-07-04 → 6 tháng đầu của sales KHÔNG có dữ liệu traffic; dùng inner join hoặc chấp nhận NaN. Ngoài ra, web_traffic KHÔNG có cột conversion_rate — phải tự tính qua orders/sessions."
    )

    doc.add_page_break()

    # ========== 4. BRAINSTORM ==========
    add_heading(doc, "4. Brainstorm — ý tưởng tổng hợp & chiến lược", 1)

    add_heading(doc, "4.1. Narrative xuyên suốt (đề xuất)", 2)
    add_callout(
        doc, "Tagline báo cáo",
        "\"Từ DATA đến DECISION — Hành trình tăng trưởng 10 năm của một thương hiệu thời trang Việt Nam.\" "
        "Cấu trúc story: Quá khứ (2013-2019: tăng trưởng) → Cú sốc COVID (2020-2021) → Phục hồi (2022) → "
        "Dự báo tương lai (2023-2024) → Khuyến nghị hành động."
    )

    add_heading(doc, "4.2. Ý tưởng EDA độc đáo (ăn điểm Sáng tạo + Insight)", 2)
    add_para(doc,
             "Mỗi ý tưởng dưới gắn tag datasets + join keys để người nhận việc biết đụng đến CSV nào. "
             "Cột 'Tầng' dùng viết tắt: D=Descriptive, Di=Diagnostic, P=Predictive, Pr=Prescriptive.")
    add_table(
        doc,
        ["#", "Ý tưởng", "Datasets & join key", "Tầng", "Output & business hook"],
        [
            ["1", "Customer lifecycle & cohort retention",
             "customers (signup_date, acquisition_channel) ⨝ orders (customer_id)",
             "Di + Pr",
             "Heatmap retention 30/60/90 × acquisition_channel → chọn kênh ROI dài hạn. "
             "Cross với geography (zip → region) để thấy cohort theo vùng."],
            ["2", "Promo ROI deep-dive (50 campaign/10.5 năm — sparse)",
             "promotions ⨝ order_items (promo_id ∥ promo_id_2) ⨝ orders (order_id, order_date) ⨝ sales (Date)",
             "Di + Pr",
             "Incremental revenue/campaign so với baseline cùng tuần năm trước; "
             "top/bottom 5 campaign; thời điểm promo nên tránh/ưu tiên."],
            ["3", "Return reason × size × category heatmap",
             "returns ⨝ order_items (order_id + product_id) ⨝ products (product_id)",
             "Di + Pr",
             "Cluster wrong_size; nếu giảm 30% wrong_size → tiết kiệm bao nhiêu refund/năm (tính từ refund_amount)? "
             "Action: fix size guide cho top SKU."],
            ["4", "Regional growth map (East/Central/West)",
             "order_items ⨝ orders (order_id) ⨝ geography (orders.zip → geography.zip → region)",
             "D + P",
             "Choropleth region × year; rising region → tăng inventory + marketing. "
             "Overlay shipments.delivery_date − ship_date để thấy logistics gap."],
            ["5", "Inventory–sales mismatch",
             "inventory (snapshot_date, product_id, stockout_days) ⨝ order_items (product_id) ⨝ orders (order_date) ⨝ sales (Date)",
             "Di + Pr",
             "Bao nhiêu lần stockout rơi đúng peak demand? Quantify lost revenue = "
             "avg_daily_qty × stockout_days × unit_price ở SKU-level."],
            ["6", "Web-traffic funnel (orders_per_session proxy)",
             "web_traffic (date) ⨝ orders (order_date); attribute theo orders.order_source vs web_traffic.traffic_source",
             "D + P",
             "Conversion proxy = n_orders/sessions theo channel; bounce_rate × AOV (từ order_items) theo source. "
             "⚠ traffic chỉ từ 2013-01 và không có cho test 2023+."],
            ["7", "Payment & device behaviour",
             "orders (payment_method, device_type, order_status) ⨝ payments (order_id, installments, payment_value)",
             "D + Di",
             "Cancel rate × payment_method (Q8 gợi ý); mobile vs desktop AOV; installments ảnh hưởng thế nào đến value."],
            ["8", "Review → return correlation",
             "reviews (rating, product_id, order_id) ⨝ orders (order_id) ⨝ returns (order_id) ⨝ products",
             "Di + P",
             "Rating ≤2 có predict return không? n-gram trên review_title cho topic (simple English). "
             "Quantify: % orders có review ≤2 bị return."],
            ["9", "Discount cannibalization",
             "promotions ⨝ order_items (promo_id) ⨝ orders (customer_id, order_date) ⨝ customers (signup_date)",
             "Di + Pr",
             "Phân biệt first-time vs returning customer trong promo window. "
             "Promo kéo khách MỚI thật không, hay chỉ giảm giá khách cũ vốn sẽ mua?"],
            ["10", "Seasonality Tết / 11.11 / 12.12 / Black Friday",
             "sales (Date) + self-built VN calendar + promotions (start_date/end_date)",
             "P + Pr",
             "Uplift % theo năm (so với baseline 14 ngày trước); dự báo tuần trước Tết 2024. "
             "Feed thẳng vào model Phần 3 như feature holiday_flag."],
            ["11", "Gross-margin × segment mix",
             "products (price, cogs, segment, category) ⨝ order_items (quantity) ⨝ orders (order_date)",
             "D + Pr",
             "Margin % theo segment (Q2); segment nào đang shrink dù margin cao → "
             "khuyến nghị marketing hoặc bundle."],
            ["12", "Shipping fee vs return rate",
             "shipments (shipping_fee) ⨝ orders ⨝ returns (order_id)",
             "Di + Pr",
             "Free-shipping có giảm return rate không? Trade-off chi phí ship vs churn."],
        ],
        widths=[0.7, 3.5, 5.0, 1.5, 6.3],
        font_size=9,
    )
    add_para(doc,
             "★ Mọi ý tag Pr phải kết nối 1 trong 3 business goals (inventory / promo / logistics — CLAUDE.md §1) "
             "khi viết insight báo cáo.",
             italic=True, color=GREY, size=10)

    add_heading(doc, "4.3. Ý tưởng Model (Phần 3) — gắn dataset nguồn feature", 2)
    add_para(doc,
             "Mục tiêu: predict Revenue + COGS hàng ngày 2023-01-01 → 2024-07-01 (548 dòng). "
             "Bảng dưới liệt kê từng khối feature + CSV nguồn + ràng buộc anti-leakage.")
    add_table(
        doc,
        ["Thành phần", "Dataset nguồn feature", "Ghi chú / ràng buộc"],
        [
            ["Baseline seasonal × growth",
             "sales (duy nhất)",
             "Geometric mean YoY growth × seasonal profile (month, day). "
             "Benchmark tối thiểu — model nâng cao phải beat ≥ 20% MAE."],
            ["Lag & rolling features",
             "sales.Revenue, sales.COGS",
             "Lag 7/14/28/365 + rolling mean/std 7/30. "
             "Test xa 548 ngày → phải recursive forecast hoặc dùng lag ≥ 548."],
            ["Calendar features",
             "Self-built (Tết ÂL, Quốc khánh 2/9, 30/4-1/5, Black Friday, 11.11, 12.12)",
             "⚠ KHÔNG import Prophet holidays — bị coi external data → loại Phần 3. "
             "Tự build DataFrame holiday từ list cố định."],
            ["Order aggregates",
             "orders group by order_date → {n_orders, n_unique_customers, pct_cancelled, pct_cod, pct_mobile}; "
             "order_items group by order_date → {n_lines, n_units_sold, pct_promo, avg_discount_pct}",
             "Aggregates này KHÔNG có cho test period 2023+ → "
             "chỉ dùng dưới dạng LAG (ví dụ lag 365) hoặc tự forecast ở mô hình phụ."],
            ["Traffic features",
             "web_traffic group by date → {sessions, bounce_rate, page_views, pct_paid_search, pct_social}",
             "Chỉ có 2013-01-01 → 2022-12-31. Test không có → quyết định 3 option (CLAUDE.md §12 #4)."],
            ["Promo indicators",
             "promotions.start_date/end_date/promo_type/applicable_category expand daily",
             "Flag active_promo_today, n_active_promos, days_since_last_promo. "
             "50 campaign đã biết lịch sử 2013-2022 → phải đoán lịch promo 2023-2024 hoặc "
             "dùng historic seasonality như proxy."],
            ["Inventory signals",
             "inventory (snapshot_date, product_id) → reindex daily forward-fill → aggregate theo ngày "
             "{total_stock_on_hand, pct_SKU_stockout, avg_days_of_supply}",
             "Snapshot end-of-month, không có cho 2023+ → forward-fill 2022-12-31 snapshot, "
             "hoặc drop feature. Decide cùng lúc với traffic."],
            ["Model chính",
             "Tất cả feature trên; 2 head riêng cho Revenue và COGS",
             "LightGBM với random_state=42, early_stopping theo temporal val set. "
             "Cân nhắc 2 model độc lập hoặc 1 model MultiOutput."],
            ["Ensemble",
             "Baseline × LightGBM (× optional SARIMAX chạy trên sales)",
             "Weighted average, weight chọn trên val set. Robust hơn single model khi COVID-like shock."],
            ["Temporal CV",
             "sales index (chỉ chia theo thời gian)",
             "Expanding window: train[...2018] → val 2019; train[...2019] → val 2020 (COVID shock!); "
             "train[...2020] → val 2021; train[...2021] → val 2022. KHÔNG random split."],
            ["Explainability (8đ báo cáo)",
             "SHAP TreeExplainer trên LightGBM + permutation importance + PDP",
             "Summary plot + 1-2 force plot ngày bất thường. "
             "Diễn giải business: 'Tết → lag 365 dominate vì pattern năm trước.'"],
        ],
        widths=[3.5, 6.3, 7.2],
        font_size=9,
    )

    add_heading(doc, "4.4. Tầng Prescriptive (nhấn mạnh — top team được đánh giá ở đây)", 2)
    add_para(doc, "Cột 'CSV nguồn & metric' liệt kê file cụ thể để phản biện khi BGK hỏi 'dựa vào đâu?'.")
    add_table(
        doc,
        ["Domain", "Khuyến nghị mẫu", "CSV nguồn & metric dẫn xuất"],
        [
            ["Pricing & discount",
             "Tăng giá segment Premium 5% (margin cao, demand ít co giãn).",
             "products (price, cogs, segment) + order_items (quantity, unit_price, discount_amount) + "
             "promotions → margin % + price-elasticity proxy khi promo."],
            ["Inventory reorder",
             "Reorder threshold = days_of_supply × 1.3 cho top SKU.",
             "inventory (stockout_flag, days_of_supply, sell_through_rate, reorder_flag) + "
             "order_items (quantity) + demand forecast từ sales."],
            ["Channel-mix budget",
             "Shift 20% từ social_media sang paid_search.",
             "web_traffic (sessions, bounce_rate, traffic_source) + orders (order_source) + "
             "order_items (revenue). ⚠ Không có cột conversion_rate — tự tính orders/sessions."],
            ["Customer retention",
             "Email re-engagement sau 90 ngày (~ median inter-order gap từ Q1).",
             "customers (signup_date, acquisition_channel) + orders (order_date, customer_id) → "
             "cohort retention × inter-order gap distribution."],
            ["Regional logistics",
             "Ưu tiên warehouse/3PL ở rising region (East vs Central vs West).",
             "orders (zip) + geography (region) + order_items (quantity, revenue) + "
             "shipments (delivery_date − ship_date = lead time)."],
            ["Size-guide fix",
             "Cập nhật size chart cho top 3 SKU có wrong_size return cao nhất.",
             "returns (return_reason='wrong_size', refund_amount) + products (size, category) + "
             "order_items → refund tiết kiệm/năm = count_wrong_size × avg_refund × giảm_% mục tiêu."],
            ["Promo scheduling",
             "Không chạy promo trong tuần Tết +7 ngày (demand đã peak tự nhiên).",
             "promotions (start_date) + sales (Date, Revenue) + self-built Tết calendar → "
             "uplift % promo vs non-promo trong và ngoài mùa Tết."],
        ],
        widths=[3.2, 6.3, 7.5],
        font_size=9,
    )

    # §4.5 — Dataset coverage matrix
    add_heading(doc, "4.5. Dataset coverage matrix (tự check trước khi submit)", 2)
    add_para(doc,
             "Ma trận 15 CSV × 4 tầng phân tích — giúp team chắc chắn đã 'đụng' đủ dataset. "
             "Ký hiệu: ✓ = bắt buộc phủ; — = không cần (file chỉ dùng cho Analytical target).")
    add_table(
        doc,
        ["CSV", "D (What)", "Di (Why)", "P (Likely)", "Pr (Action)"],
        [
            ["products", "✓", "✓", "—", "✓"],
            ["customers", "✓", "✓", "✓", "✓"],
            ["geography", "✓", "✓", "—", "✓"],
            ["promotions", "—", "✓", "✓", "✓"],
            ["orders", "✓", "✓", "✓", "✓"],
            ["order_items", "✓", "✓", "✓", "✓"],
            ["payments", "✓", "✓", "—", "—"],
            ["shipments", "✓", "✓", "—", "✓"],
            ["returns", "✓", "✓", "—", "✓"],
            ["reviews", "✓", "✓", "—", "—"],
            ["inventory", "✓", "✓", "✓", "✓"],
            ["web_traffic", "✓", "✓", "✓", "✓"],
            ["sales", "✓", "—", "✓", "✓"],
            ["sample_submission", "—", "—", "✓", "—"],
        ],
        widths=[3.5, 3.3, 3.3, 3.3, 3.6],
        font_size=9,
    )

    doc.add_page_break()

    # ========== 5. PLAN 13 NGÀY ==========
    add_heading(doc, "5. Plan 13 ngày → deadline 2026-05-01", 1)
    add_table(
        doc,
        ["Giai đoạn", "Ngày", "Output", "Lead"],
        [
            ["Kickoff + Data audit",
             "04-18 → 04-19",
             "01_data_audit.ipynb (validate schema, missing, outliers, integrity check)",
             "All / Lead Analyst"],
            ["MCQ",
             "04-20",
             "02_mcq.ipynb + bộ 10 đáp án (verify chéo)",
             "Hiển"],
            ["EDA Descriptive + Diagnostic",
             "04-21 → 04-23",
             "03_eda_descriptive.ipynb — tầng 1-2, 6-8 charts",
             "Lead Analyst"],
            ["EDA Predictive + Prescriptive",
             "04-24 → 04-26",
             "04_eda_predictive.ipynb — tầng 3-4, story-led, action-oriented",
             "Insights Lead"],
            ["Feature engineering + Model",
             "04-27 → 04-29",
             "05_FE.ipynb, 06_baseline.ipynb, 07_advanced.ipynb + Kaggle submission",
             "ML Engineer"],
            ["Report + GitHub + form",
             "04-30 → 05-01",
             "main.pdf ≤4 trang (NeurIPS), README.md, submit form chính thức",
             "Hiển + All"],
        ],
        widths=[4.5, 3.0, 7.0, 2.5],
    )
    add_callout(
        doc, "Dự phòng",
        "Buffer 1 ngày (05-01) cho việc submit cuối + fix lỗi format PDF/Kaggle. "
        "Không lùi deadline nội bộ sang 04-30 trừ khi có rủi ro lớn."
    )
    add_callout(
        doc, "Rebuild baseline",
        "baseline.ipynb hiện dùng DATA_DIR='dataset/' — sai path so với folder "
        "dataset-datathon-2026-round-1/. Trong giai đoạn 04-27 → 04-29 (ML Engineer), "
        "bước đầu tiên là fork baseline sang 06_model_baseline.ipynb, sửa DATA_DIR, re-run end-to-end "
        "để xác nhận reproducibility trước khi build LightGBM.",
        color=NAVY,
    )

    # ========== 6. PHÂN CÔNG ==========
    add_heading(doc, "6. Phân công vai trò (đề xuất)", 1)
    add_para(doc, "Hiển đã chốt làm đội trưởng. 3 thành viên còn lại cần điền vào 3 role sau "
             "(quyết định trong cuộc họp đầu):")
    add_table(
        doc,
        ["Role", "Trách nhiệm", "Notebook chính", "Người"],
        [
            ["Đội trưởng",
             "Coordinator, MCQ, LaTeX báo cáo, GitHub, submission form, theo deadline.",
             "02, report/, README",
             "Hiển ✅"],
            ["Lead Analyst",
             "EDA tầng 1-2 (Descriptive + Diagnostic), data audit, storytelling.",
             "01, 03",
             "(Phúc / Kiên / Đồng)"],
            ["Insights Lead",
             "EDA tầng 3-4 (Predictive + Prescriptive), viết insight cho báo cáo.",
             "04",
             "(Phúc / Kiên / Đồng)"],
            ["ML Engineer",
             "Feature engineering, baseline, LightGBM, SHAP, Kaggle submissions.",
             "05, 06, 07",
             "(Phúc / Kiên / Đồng)"],
        ],
        widths=[3.0, 7.0, 4.0, 3.0],
    )
    add_para(doc, "Tất cả thành viên: review chéo notebook người khác, đảm bảo seed=42 và chạy được "
             "trên máy khác.", italic=True, color=GREY)

    # ========== 7. RỦI RO ==========
    add_heading(doc, "7. Rủi ro & mitigation", 1)
    add_table(
        doc,
        ["Rủi ro", "Tác động", "Mitigation"],
        [
            ["RAM laptop không đủ cho orders.csv (44MB) + order_items.csv (23MB)",
             "Notebook crash khi load full",
             "Khai báo dtype, dùng chunksize, hoặc move sang Kaggle/Colab."],
            ["Prophet auto-load Vietnam holiday → bị coi 'external data'",
             "Loại Phần 3 (mất 20đ)",
             "KHÔNG dùng Prophet với holiday. Tự xây calendar VN."],
            ["Dùng Revenue/COGS 2023+ làm feature (leakage)",
             "Loại Phần 3",
             "Anti-leakage checklist trong CLAUDE.md §9. Code review cross check."],
            ["web_traffic chỉ có đến 2022-12-31, test 2023-2024 không có",
             "Mất feature mạnh",
             "(a) Forecast traffic riêng / (b) Rolling mean 2022 làm proxy / (c) Bỏ traffic feature. "
             "→ Quyết định ở Phase 6."],
            ["PDF báo cáo ≤4 trang RẤT chặt cho cả EDA + Model",
             "Cắt insight",
             "Plan layout từ tuần 1, ưu tiên 4-5 chart đắt giá nhất + bảng tóm tắt."],
            ["Repository public → ảnh thẻ SV bị leak nếu commit nhầm",
             "Vi phạm bảo mật cá nhân",
             "Thêm photos/, .env vào .gitignore. KHÔNG push ảnh thẻ. Upload ảnh thẻ riêng qua form."],
            ["Screen recording .mov 5GB",
             "Push fail",
             "Thêm *.mov vào .gitignore."],
        ],
        widths=[5.5, 4.0, 7.5],
    )

    # ========== 8. CHECKLIST ==========
    add_heading(doc, "8. Checklist nộp bài", 1)
    add_para(doc, "Tick từng mục trước khi bấm Submit form:", bold=True)
    add_bullets(doc, [
        "GitHub repo public hoặc invite organizers; README.md có hướng dẫn reproduce đầy đủ.",
        "Kaggle submission đúng 548 dòng, thứ tự khớp 100% sample_submission.csv.",
        "Báo cáo PDF ≤4 trang (không tính refs/appendix), template NeurIPS, có link GitHub.",
        "Báo cáo có mục Explainability (SHAP / feature importance / PDP) bằng ngôn ngữ business.",
        "Form: đáp án 10 MCQ + PDF + GitHub link + Kaggle link + ảnh thẻ SV 4 thành viên.",
        "Tickbox xác nhận ≥1 thành viên tham dự Vòng Chung kết 2026-05-23 tại VinUni HN.",
        "Random seed = 42 set consistent ở mọi notebook.",
        "Không có Revenue/COGS test set xuất hiện trong feature pipeline.",
        "Không dùng dữ liệu/library auto-load external data (Prophet holiday, etc.).",
    ])

    # ========== 9. CÂU HỎI MỞ ==========
    add_heading(doc, "9. Câu hỏi mở cho cuộc họp đầu tiên", 1)
    add_para(doc, "Cần thống nhất trong cuộc họp team đầu tiên:", bold=True)
    add_bullets(doc, [
        ("GitHub repo",
         "Tên repo? Owner account? Public từ đầu hay private rồi public sát deadline?"),
        ("Kaggle team setup",
         "Tất cả 4 thành viên đã join competition chưa? Merge team trên Kaggle hay 1 account submit?"),
        ("Môi trường compute",
         "Ai dùng laptop? RAM bao nhiêu? Có cần Colab/Kaggle Notebook làm fallback không?"),
        ("Web traffic test period",
         "Chọn option (a) self-forecast / (b) rolling mean proxy / (c) bỏ feature?"),
        ("Inventory test period",
         "inventory snapshot dừng 2022-12-31, test 18 tháng 2023-01 → 2024-07 không có. "
         "Forward-fill snapshot 2022-12 hay drop feature? (Nên quyết cùng lúc với web_traffic.)"),
        ("Prophet vs LightGBM",
         "Có dùng Prophet không (rủi ro external data)? Hay LightGBM-only ngay từ đầu?"),
        ("Phân công 3 role còn lại",
         "Phúc / Kiên / Đồng nhận role nào (Lead Analyst / Insights Lead / ML Engineer)?"),
        ("Lịch họp",
         "Họp chốt mỗi mấy ngày (đề xuất: 04-19, 04-23, 04-27, 04-30)?"),
    ])

    # ========== APPENDIX ==========
    add_heading(doc, "Phụ lục — File tham chiếu nội bộ", 1)
    add_bullets(doc, [
        ("CLAUDE.md", "Brain file tổng cho Claude sessions — đọc trước khi làm việc."),
        ("Đề thi Vòng 1.pdf", "Đề chính thức — luôn là nguồn cuối cùng nếu có tranh cãi."),
        ("Rule_book.docx", "Quy chế chi tiết, ít chi tiết hơn PDF nhưng có business context."),
        ("dataset-datathon-2026-round-1/baseline.ipynb", "Baseline có sẵn, dùng làm benchmark tối thiểu."),
        ("dataset-datathon-2026-round-1/dataset_description.docx", "Data dictionary chi tiết."),
    ])

    add_para(doc,
             "Tài liệu này là snapshot tại 2026-04-18. Sau cuộc họp đầu, update lại §6 (phân công), "
             "§9 (resolved questions), và lưu thành phiên bản v2 nếu cần.",
             italic=True, color=GREY, size=10, space_after=0)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    out = build()
    size_kb = out.stat().st_size / 1024
    print(f"Wrote: {out}")
    print(f"Size:  {size_kb:.1f} KB")
